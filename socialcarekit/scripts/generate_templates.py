#!/usr/bin/env python3
"""Generate the SocialCareKit downloadable template library.

Creates properly formatted, genuinely usable DOCX templates (python-docx) and
one XLSX tracker (openpyxl) into storage/templates/files/, plus the
database/seed-content/templates.json seed file.

Run:  python3 scripts/generate_templates.py
"""

import json
import os
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUT_DIR = os.path.join(BASE, "storage", "templates", "files")
SEED_DIR = os.path.join(BASE, "database", "seed-content")

INTRO = ("Part of the SocialCareKit template library — socialcarekit.com. "
         "Free for use within your organisation. Review before use.")
FOOTER = "Template last reviewed: July 2026"

BRAND = RGBColor(0x1F, 0x4E, 0x5F)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_FILL = "1F4E5F"
LIGHT_FILL = "E9F1F3"

PORTRAIT_W = 17.4   # usable content width in cm (A4 minus margins)
LANDSCAPE_W = 26.1

CHECK = "☐"  # empty checkbox


# ---------------------------------------------------------------- helpers ---

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def fixed_layout(table):
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.insert_element_before(
        layout, "w:tblCellMar", "w:tblLook", "w:tblCaption",
        "w:tblDescription")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def row_height(row, cm):
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    row.height = Cm(cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def cell_text(cell, text="", bold=False, italic=False, size=None, color=None,
              fill=None, align=None):
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def new_doc(title, landscape=False):
    doc = Document()
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    else:
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for name, size in (("Title", 20), ("Heading 1", 14), ("Heading 2", 12),
                       ("Heading 3", 11)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BRAND

    # Fix python-docx default template quirk: w:zoom requires w:percent
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    doc.add_paragraph(title, style="Title")
    p = doc.add_paragraph()
    r = p.add_run(INTRO)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return doc


def finish(doc, slug):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(FOOTER)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "999999")
    p_bdr.append(top)
    p_pr.insert_element_before(
        p_bdr, "w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr",
        "w:sectPr")
    path = os.path.join(OUT_DIR, slug + ".docx")
    doc.save(path)
    return path


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text, italic=False, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def guidance(doc, text):
    return para(doc, text, italic=True, size=9, color=GREY)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def field_table(doc, fields, header=None, label_w=6.0, row_h=0.85,
                total_w=PORTRAIT_W):
    """Two-column label + generous blank cell table.

    fields: list of strings, or (label, hint) tuples where the hint is
    printed in the blank cell in small grey italics.
    """
    value_w = total_w - label_w
    n_header = 1 if header else 0
    table = doc.add_table(rows=n_header + len(fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(table)
    idx = 0
    if header:
        row = table.rows[0]
        row.cells[0].merge(row.cells[1])
        cell_text(row.cells[0], header, bold=True, color=WHITE, fill=DARK_FILL)
        repeat_header(row)
        row_height(row, 0.6)
        idx = 1
    for f in fields:
        label, hint = (f, "") if isinstance(f, str) else f
        row = table.rows[idx]
        row_height(row, row_h)
        cell_text(row.cells[0], label, bold=True, fill=LIGHT_FILL)
        if hint:
            cell_text(row.cells[1], hint, italic=True, size=8.5, color=GREY)
        row.cells[0].width = Cm(label_w)
        row.cells[1].width = Cm(value_w)
        idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def grid_table(doc, headers, widths, blank_rows=6, row_h=1.0, font_size=9.5,
               prefill=None):
    """Multi-column table with a shaded header row and blank data rows.

    prefill: optional list of rows (list of strings) written before the
    blank rows.
    """
    prefill = prefill or []
    table = doc.add_table(rows=1 + len(prefill) + blank_rows, cols=len(headers))
    table.style = "Table Grid"
    fixed_layout(table)
    hrow = table.rows[0]
    repeat_header(hrow)
    row_height(hrow, 0.6)
    for i, htext in enumerate(headers):
        cell_text(hrow.cells[i], htext, bold=True, color=WHITE,
                  fill=DARK_FILL, size=font_size)
        hrow.cells[i].width = Cm(widths[i])
    r_idx = 1
    for values in prefill:
        row = table.rows[r_idx]
        row_height(row, 0.55)
        for i, v in enumerate(values):
            cell_text(row.cells[i], v, size=font_size)
            row.cells[i].width = Cm(widths[i])
        r_idx += 1
    for _ in range(blank_rows):
        row = table.rows[r_idx]
        row_height(row, row_h)
        for i in range(len(headers)):
            row.cells[i].width = Cm(widths[i])
        r_idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def box(doc, label, height=3.0, hint=""):
    """Section label bar with a tall blank writing box beneath it."""
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    fixed_layout(table)
    cell_text(table.rows[0].cells[0], label, bold=True, fill=LIGHT_FILL)
    row_height(table.rows[0], 0.6)
    row_height(table.rows[1], height)
    if hint:
        cell_text(table.rows[1].cells[0], hint, italic=True, size=8.5,
                  color=GREY)
    table.rows[0].cells[0].width = Cm(PORTRAIT_W)
    table.rows[1].cells[0].width = Cm(PORTRAIT_W)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def check_table(doc, items, first_col="Item", extra=("Date", "Initials"),
                first_w=None, extra_w=2.6, total_w=PORTRAIT_W):
    if first_w is None:
        first_w = total_w - extra_w * len(extra)
    headers = [first_col] + list(extra)
    widths = [first_w] + [extra_w] * len(extra)
    return grid_table(doc, headers, widths, blank_rows=0, prefill=[
        [it] + [""] * len(extra) for it in items])


def sig_table(doc, roles, total_w=PORTRAIT_W):
    widths = [total_w - 4.5 - 5.0 - 3.2, 4.5, 5.0, 3.2]
    return grid_table(doc, ["Role", "Name", "Signature", "Date"], widths,
                      blank_rows=0, row_h=0.9,
                      prefill=[[r, "", "", ""] for r in roles])


def yn_line(items):
    return "   ".join("{} {}".format(CHECK, i) for i in items)


# ---------------------------------------------------------- OFSTED builders ---

def t_placement_plan():
    doc = new_doc("Placement Plan")
    guidance(doc, "Complete within the first 5 working days of placement and "
                  "review alongside the child's care plan and statutory "
                  "reviews. This plan records how the home will meet the "
                  "child's day-to-day needs under Regulation 5 of the Care "
                  "Planning, Placement and Case Review (England) Regulations "
                  "2010.")
    h1(doc, "1. Child's details")
    field_table(doc, [
        "Full name", "Preferred name", "Date of birth", "Age at placement",
        "Gender / pronouns", "Ethnicity", "Religion", "First language",
        ("Communication needs", "e.g. Makaton, interpreter, visual aids"),
        "NHS number", "Date placement commenced",
        "Planned length / end date of placement",
    ], header="Child's details")

    h1(doc, "2. Legal status and placing authority")
    field_table(doc, [
        ("Legal status", "e.g. section 20 CA 1989, interim care order, full "
                         "care order, remand"),
        "Who holds parental responsibility",
        "Placing authority", "Allocated social worker (name, phone, email)",
        "Out-of-hours / EDT number",
        "Independent Reviewing Officer (IRO) and contact details",
        "Virtual School contact", "Date of next LAC / statutory review",
    ], header="Placing authority and key contacts")

    h1(doc, "3. Delegated authority and consents")
    guidance(doc, "Record who may consent to each matter and any conditions. "
                  "Attach the placing authority's delegated authority "
                  "agreement where available.")
    grid_table(doc, ["Matter", "Consent held by (home / SW / parent)",
                     "Conditions / notes"],
               [6.4, 5.0, 6.0], blank_rows=0, prefill=[
        ["Routine medical and dental treatment", "", ""],
        ["Emergency medical treatment", "", ""],
        ["Immunisations", "", ""],
        ["Haircuts and personal appearance", "", ""],
        ["School trips and activities", "", ""],
        ["Overnight stays with friends", "", ""],
        ["Photographs (in-home use / external)", "", ""],
        ["Holidays (UK / abroad, passport)", "", ""],
        ["Other (specify)", "", ""],
    ])

    h1(doc, "4. Health")
    field_table(doc, [
        "GP (name, address, phone)", "Dentist", "Optician",
        "Date of last Initial / Review Health Assessment",
        ("Current medication", "name, dose, times, who administers"),
        "Allergies", "Known health conditions",
        ("Emotional wellbeing and mental health support",
         "e.g. CAMHS worker, therapy, waiting lists"),
        "Health promotion needs (diet, exercise, smoking, sexual health)",
    ], header="Health arrangements")

    h1(doc, "5. Education")
    field_table(doc, [
        "School / education provision and year group",
        "Designated teacher (name and contact)",
        "Date of current Personal Education Plan (PEP)",
        ("SEND / EHCP", "status, date of plan, key provisions"),
        "Attendance arrangements and transport",
        "Homework, equipment and study support in the home",
        "Aspirations and post-16 planning",
    ], header="Education arrangements")

    h1(doc, "6. Identity, culture and religion")
    box(doc, "How the home will promote the child's identity, culture, "
             "religion and heritage", 3.0,
        "Include diet, worship, festivals, hair and skin care, language, "
        "community links and life-story work.")

    h1(doc, "7. Family time (contact) arrangements")
    grid_table(doc, ["Person", "Relationship", "Type / frequency",
                     "Supervised?", "Transport / venue", "Restrictions"],
               [3.2, 2.6, 3.2, 2.2, 3.2, 3.0], blank_rows=5, row_h=0.9)

    h1(doc, "8. Risks and safety planning")
    guidance(doc, "Summarise key risks; cross-reference the child's individual "
                  "risk assessments, behaviour support plan and any missing "
                  "from care protocol.")
    grid_table(doc, ["Risk", "Level (L/M/H)", "How it will be managed",
                     "Linked plan / assessment"],
               [4.4, 2.0, 7.0, 4.0], blank_rows=5, row_h=1.0)

    h1(doc, "9. Daily living")
    field_table(doc, [
        "Bedtime and morning routine", "Pocket money amount and day",
        "Savings arrangements", "Mobile phone / internet agreement",
        "Chores and independence tasks", "Hobbies, clubs and activities",
        "Dietary needs and preferences",
    ], header="Daily routines")

    h1(doc, "10. Permanence and moving on")
    box(doc, "Plan for permanence / transition from this placement", 2.5,
        "Long-term plan, planned moves, preparation for independence where "
        "relevant.")

    h1(doc, "11. Agreement and signatures")
    sig_table(doc, ["Registered manager", "Key worker",
                    "Placing authority social worker",
                    "Child / young person (where appropriate)",
                    "Parent / person with PR (where appropriate)"])
    field_table(doc, ["Date plan completed", "Date of next review"])
    return doc


def t_behaviour_support_plan():
    doc = new_doc("Behaviour Support Plan (PBS format)")
    guidance(doc, "A positive behaviour support (PBS) plan. Complete with the "
                  "child wherever possible, share with the whole team, and "
                  "review after any significant incident and at least "
                  "monthly.")
    h1(doc, "1. About the young person")
    field_table(doc, [
        "Name", "Date of birth", "Date of plan", "Plan written by",
        "Contributors (child, family, SW, CAMHS, school)",
        "Review date",
    ], header="Plan details")
    box(doc, "Strengths, interests and what a good day looks like", 2.5)

    h1(doc, "2. Understanding the behaviour")
    box(doc, "What does the behaviour look like? (describe it factually)",
        2.5, "Describe observable behaviour, not labels — e.g. 'shouts and "
             "kicks doors' rather than 'kicks off'.")
    box(doc, "Triggers and setting events", 2.5,
        "e.g. family time, transitions, noise, hunger, anniversaries, peer "
        "dynamics, feeling out of control.")
    box(doc, "What is the behaviour communicating? (its function)", 2.5,
        "e.g. seeking safety, connection or control; escaping demands; "
        "sensory needs; expression of trauma.")
    box(doc, "Early warning signs that things are building", 2.0,
        "Changes in body language, tone, withdrawal, pacing, etc.")

    h1(doc, "3. Primary (proactive) strategies — keeping things calm")
    box(doc, "What we do every day to reduce the likelihood of incidents",
        3.0, "Routine, relationships, environment, communication approach, "
             "activities, sensory strategies, choices and control.")

    h1(doc, "4. Secondary (de-escalation) strategies — when signs appear")
    box(doc, "What helps when the young person is becoming distressed", 3.0,
        "Named adults, scripts and phrases that help, space, distraction, "
        "reduced demands. Include what does NOT help.")

    h1(doc, "5. Reactive strategies — during an incident")
    box(doc, "Agreed responses to keep everyone safe", 2.5,
        "Include any authorised physical intervention holds from your "
        "accredited training model, who may use them, and last-resort "
        "criteria. Restraint only to prevent injury, serious damage to "
        "property or absconding risk of harm (Reg 20, Children's Homes "
        "(England) Regulations 2015).")
    field_table(doc, [
        ("Physical intervention authorised?", "Yes / No — specify techniques"),
        "Medical considerations relevant to restraint",
        "Positions / holds that must NOT be used",
    ])

    h1(doc, "6. After an incident — repair and reflection")
    box(doc, "Post-incident support for the young person and staff", 2.5,
        "Debrief with the child (within 24–48 hours), restorative approach, "
        "staff debrief, recording (daily log, physical intervention record), "
        "notifications.")

    h1(doc, "7. Agreement")
    sig_table(doc, ["Young person (if able / willing)", "Key worker",
                    "Registered manager"])

    doc.add_page_break()
    h1(doc, "ABC recording sheet")
    guidance(doc, "Complete a row after each incident of concern. Use the "
                  "pattern over time to review sections 2–5 of this plan.")
    field_table(doc, ["Young person", "Month"], label_w=4.5)
    grid_table(doc, ["Date / time & staff initials",
                     "Antecedent — what happened just before?",
                     "Behaviour — what exactly did they do?",
                     "Consequence — what happened next / how did it end?"],
               [3.2, 4.8, 4.7, 4.7], blank_rows=7, row_h=1.7, font_size=9)
    return doc


def t_physical_intervention_record():
    doc = new_doc("Physical Intervention Record")
    guidance(doc, "Complete within 24 hours of any use of a measure of "
                  "control, discipline or restraint. Regulation 35 of the "
                  "Children's Homes (England) Regulations 2015 requires the "
                  "record to be completed within 48 hours, the child to be "
                  "given the opportunity to add their view within 5 days, and "
                  "the registered person to add a signed note confirming "
                  "review within 48 hours of the record being made.")
    h1(doc, "1. Basic details")
    field_table(doc, [
        "Name of child", "Date of incident", "Time incident began / ended",
        "Location", "Record completed by (name and role)",
        "Date and time record completed",
    ], header="Incident details")

    h1(doc, "2. People involved")
    grid_table(doc, ["Name", "Role (staff / child / witness)",
                     "Involvement (restrained, held, observed, supported)"],
               [4.6, 4.4, 8.4], blank_rows=4, row_h=0.7)

    h1(doc, "3. What led up to the incident")
    box(doc, "Antecedents — what was happening beforehand?", 2.5)
    box(doc, "De-escalation attempted before physical intervention", 2.5,
        "List every strategy tried (verbal, space, distraction, change of "
        "face) and the child's response.")

    h1(doc, "4. The intervention")
    field_table(doc, [
        ("Behaviour that made intervention necessary",
         "Describe factually — risk of injury, serious property damage, "
         "absconding into danger"),
        ("Technique(s) used", "Name of hold per your accredited training "
                              "model (e.g. Team Teach, PRICE)"),
        "Number of staff involved in the hold", "Duration of the hold",
        "How the intervention ended",
    ], header="Details of the physical intervention", row_h=1.1)

    h1(doc, "5. Injuries and medical attention")
    field_table(doc, [
        ("Injuries to the child", "Include 'none apparent' if so; complete "
                                  "body map if injured"),
        "Injuries to staff or others",
        "First aid / medical attention given (by whom, when)",
        "GP / hospital attendance required? (details)",
    ], row_h=1.0)

    h1(doc, "6. Effectiveness and consequences")
    box(doc, "Effectiveness of the measure and any consequences for the "
             "child", 2.2)

    h1(doc, "7. Child's views")
    box(doc, "The child's account and views of the incident (within 5 days)",
        2.8, "Record verbatim where possible. Note who supported the child, "
             "when, and whether they declined to comment.")
    field_table(doc, ["Child's views recorded by", "Date views recorded",
                      ("Did the child add to / sign this record?", "Yes / No")])

    h1(doc, "8. Debrief and follow-up")
    field_table(doc, [
        "Debrief with child completed (date, by whom)",
        "Debrief with staff completed (date, by whom)",
        "Behaviour support plan / risk assessment reviewed? (date)",
        ("Notifications made", "Parents / PR, social worker, IRO, placing "
                               "authority; Ofsted if notifiable under Reg 40"),
    ], row_h=1.0)

    h1(doc, "9. Signatures and management review")
    sig_table(doc, ["Staff member completing record",
                    "Other staff involved",
                    "Registered manager (review within 48 hours)"])
    box(doc, "Registered person's signed review note (Reg 35(3)(b))", 2.2,
        "Comment on whether the intervention was necessary, proportionate "
        "and carried out correctly, and any learning or action required.")
    return doc


def t_missing_from_care_record():
    doc = new_doc("Missing from Care Record & Return Home Interview")
    guidance(doc, "Use Part A while the child is missing and on return. "
                  "Part B is for the independent return home interview (RHI), "
                  "which should be offered within 72 hours of return "
                  "(Statutory guidance on children who run away or go missing "
                  "from home or care, 2014).")
    h1(doc, "Part A — Missing episode record")
    h2(doc, "A1. The child and the episode")
    field_table(doc, [
        "Name of child", "Date of birth",
        "Date and time last seen / episode began", "Where last seen",
        ("Classification", "Missing / Away from placement without "
                           "authorisation"),
        ("Description and clothing", "Height, build, hair, distinguishing "
                                     "features, what they were wearing"),
        "Money, phone, bank card with them?",
    ], header="Episode details")

    h2(doc, "A2. Risk information")
    field_table(doc, [
        ("Risk level per missing risk assessment", "Low / Medium / High"),
        ("Known risks", "CSE / CCE, county lines, self-harm, substances, "
                        "risky associates, previous patterns"),
        "Likely locations and associates",
        "Grab pack / recent photo available?",
    ], row_h=1.0)

    h2(doc, "A3. Actions taken")
    grid_table(doc, ["Date / time", "Action taken", "By whom", "Outcome"],
               [2.6, 7.0, 3.0, 4.8], blank_rows=6, row_h=0.7)
    field_table(doc, [
        "Police notified (date, time, by whom)",
        "Police reference / CAD number",
        "Placing authority social worker / EDT notified (date, time)",
        "Parents / those with PR notified (where appropriate)",
        ("Ofsted notified?", "Required for serious incidents under Reg 40"),
    ])

    h2(doc, "A4. Return")
    field_table(doc, [
        "Date and time returned", "How the child returned (found by whom)",
        ("Condition on return", "Physical presentation, injuries, "
                                "intoxication, new belongings, emotional "
                                "state"),
        "Immediate needs met (food, sleep, first aid, welcome back)",
        "Police safe and well / prevention interview (date, officer)",
    ], header="Return details", row_h=1.0)

    doc.add_page_break()
    h1(doc, "Part B — Return Home Interview (independent person)")
    field_table(doc, [
        "Interview carried out by (name, role, organisation)",
        "Relationship to child (must be independent of the home)",
        "Date and time of interview (within 72 hours of return)",
        "Location", "Did the child engage? (if declined, record how offered)",
    ], header="Interview details")
    h2(doc, "Suggested questions — record answers in the child's own words")
    for q in [
        "What happened before you left? Was there anything that made you "
        "want to go?",
        "Where did you go and how did you get there?",
        "Who were you with? Did you meet anyone new?",
        "Where did you stay and how did you get food, money or lifts?",
        "Did anything happen while you were away that scared or hurt you?",
        "Did anyone ask you to do anything you didn't want to do?",
        "What made you decide to come back?",
        "What would make it less likely you'd go missing again?",
        "Is there anything you want to happen now, or anyone you want to "
        "talk to?",
    ]:
        box(doc, q, 1.4)
    h2(doc, "Outcome of interview")
    box(doc, "Concerns identified (exploitation, harm, push/pull factors)",
        2.2)
    grid_table(doc, ["Action agreed", "Owner", "By when"],
               [10.0, 4.0, 3.4], blank_rows=4, row_h=0.7)
    field_table(doc, [
        "Shared with (social worker, home manager, police where needed)",
        "Missing risk assessment updated? (date)",
    ])
    sig_table(doc, ["Independent interviewer", "Registered manager"])
    return doc


def t_key_work_session_log():
    doc = new_doc("Key Work Session Log")
    guidance(doc, "Record each planned one-to-one session. Link the session "
                  "to the child's relevant plans so progress can be evidenced "
                  "at reviews and in the Reg 45 quality of care review.")
    field_table(doc, [
        "Child / young person", "Key worker", "Date and time", "Duration",
        "Location", "Others present",
    ], header="Session details")
    field_table(doc, [
        ("Planned topic(s)", "e.g. keeping safe online, family time, "
                             "independence skills, health, feelings"),
        ("Link to care plan / placement plan objective",
         "Which target or outcome does this session support?"),
    ], row_h=1.0)
    box(doc, "What we did and talked about", 4.5,
        "Summarise the session factually. Note the child's engagement.")
    box(doc, "The child's views, wishes and feelings", 3.0,
        "Use the child's own words where possible.")
    grid_table(doc, ["Action agreed", "Who", "By when", "Done?"],
               [9.4, 3.0, 2.8, 2.2], blank_rows=4, row_h=0.7)
    field_table(doc, [
        "Anything to hand over / escalate (safeguarding, wellbeing)?",
        "Focus for next session and planned date",
    ], row_h=1.0)
    sig_table(doc, ["Key worker", "Child (optional)"])
    return doc


def t_reg_44_visit_pack():
    doc = new_doc("Regulation 44 Independent Visitor — Preparation Pack")
    guidance(doc, "For the independent person carrying out monthly visits "
                  "under Regulation 44 of the Children's Homes (England) "
                  "Regulations 2015. The visitor must form an opinion on "
                  "whether children are effectively safeguarded and whether "
                  "the conduct of the home promotes children's wellbeing. At "
                  "least half of visits each year should be unannounced. The "
                  "report should be sent to Ofsted, the registered person and "
                  "the responsible individual as soon as possible.")

    h1(doc, "1. Visit details")
    field_table(doc, [
        "Home visited", "Date and time of visit",
        ("Announced or unannounced?", "Record running total of unannounced "
                                      "visits this year"),
        "Independent person (name, organisation)",
        "Children resident on the day / children seen",
        "Staff on duty / staff spoken to",
    ], header="Visit details")

    h1(doc, "2. Document sampling checklist")
    guidance(doc, "Sample records since the last visit. Tick and note "
                  "findings against each.")
    check_table(doc, [
        "Daily logs (sample across children)",
        "Sanctions / consequences record (Reg 35)",
        "Physical intervention records (Reg 35)",
        "Missing from care episodes and return home interviews",
        "Complaints record and outcomes",
        "Safeguarding concerns and referrals",
        "Accident and incident records",
        "Medication administration records and errors",
        "Children's case files / key work sessions (sample)",
        "Staff rotas — actual staffing vs planned, agency use",
        "Supervision and team meeting records",
        "Training matrix (safeguarding, first aid, restraint refreshers)",
        "Recruitment files / single central record (sample)",
        "Fire safety checks, drills and premises / maintenance log",
        "Location risk assessment (last review date)",
        "Previous Reg 44 reports — progress on recommendations",
    ], first_col="Record sampled", extra=("Seen", "Notes / concerns"),
        extra_w=2.2)

    h1(doc, "3. Suggested questions for children")
    bullets(doc, [
        "Do you feel safe here? Is there anything or anyone that worries "
        "you?",
        "Who would you talk to if you were unhappy — and what happened last "
        "time you raised something?",
        "What are the staff like? Who do you get on with best?",
        "What's the food like? Do you get a say in meals and activities?",
        "Are consequences fair? What happens when things go wrong?",
        "How is school going? Do staff help with homework and appointments?",
        "How is time with your family and friends working?",
        "Do you know how to complain, and about your advocate?",
        "If you could change one thing about living here, what would it be?",
    ])

    h1(doc, "4. Suggested questions for staff")
    bullets(doc, [
        "When was your last supervision, and do you find it useful?",
        "What would you do if a child disclosed abuse? Who is the DSL?",
        "How would you raise a concern about a colleague or manager "
        "(whistleblowing)?",
        "Are staffing levels sufficient across the rota? How often is agency "
        "used?",
        "What training have you completed recently? Anything outstanding?",
        "Do you know each child's plans and current risks?",
        "What has changed for the better here in the last three months?",
    ])

    h1(doc, "5. Suggested questions for the manager")
    bullets(doc, [
        "What significant events / notifications have there been since the "
        "last visit?",
        "Current staffing position: vacancies, sickness, agency reliance, "
        "recruitment in progress?",
        "Progress against recommendations from my last report?",
        "Any placements made or ended? Any placements outside the home's "
        "statement of purpose?",
        "Any current safeguarding concerns, missing patterns or "
        "exploitation risks?",
        "What is keeping you awake at night about this home?",
    ])

    h1(doc, "6. Observation prompts")
    bullets(doc, [
        "Condition of the home: homely, clean, well maintained, "
        "personalised bedrooms (view with the child's consent).",
        "Atmosphere and interactions between staff and children.",
        "Food available; mealtime arrangements.",
        "Health and safety hazards; fire exits clear.",
        "Evidence of children's voice: menus, activity planning, house "
        "meetings.",
    ])

    h1(doc, "7. Report structure")
    guidance(doc, "Structure the written report as follows.")
    for label, height in [
        ("Summary of the visit and methodology", 1.8),
        ("Opinion: are children effectively safeguarded? (with evidence)",
         2.5),
        ("Opinion: does the conduct of the home promote children's "
         "wellbeing? (with evidence)", 2.5),
        ("Progress on previous recommendations", 1.8),
    ]:
        box(doc, label, height)
    grid_table(doc, ["No.", "Recommendation", "For", "Timescale"],
               [1.2, 9.4, 3.4, 3.4], blank_rows=4, row_h=0.7)
    field_table(doc, [
        "Report sent to Ofsted, registered person and responsible "
        "individual (dates)",
    ], row_h=0.9)
    sig_table(doc, ["Independent person"])
    return doc


def t_reg_45_quality_review():
    doc = new_doc("Regulation 45 Quality of Care Review")
    guidance(doc, "Complete at least once every 6 months (Reg 45, Children's "
                  "Homes (England) Regulations 2015). The review must cover "
                  "the quality of care provided (45(2)(a)), the feedback and "
                  "opinions of children (45(2)(b)) and actions taken "
                  "following complaints, and be used to improve the care "
                  "provided. Supply the written report to Ofsted within 28 "
                  "days of a request.")
    field_table(doc, [
        "Home", "Period covered (from / to)", "Completed by", "Date",
        "Children resident during the period",
    ], header="Review details")

    h1(doc, "1. Quality of care provided — Reg 45(2)(a)")
    h2(doc, "1.1 Data for the period")
    grid_table(doc, ["Measure", "This period", "Last period", "Commentary"],
               [6.2, 2.4, 2.4, 6.4], blank_rows=0, prefill=[
        ["Admissions / discharges", "", "", ""],
        ["Missing episodes (and children involved)", "", "", ""],
        ["Physical interventions", "", "", ""],
        ["Sanctions / consequences", "", "", ""],
        ["Safeguarding concerns / referrals", "", "", ""],
        ["Allegations against staff (LADO)", "", "", ""],
        ["Complaints received", "", "", ""],
        ["Accidents / incidents", "", "", ""],
        ["Medication errors", "", "", ""],
        ["Notifications to Ofsted (Reg 40)", "", "", ""],
        ["School attendance / exclusions", "", "", ""],
        ["Health appointments completed / missed", "", "", ""],
        ["Staff turnover / vacancies / agency shifts", "", "", ""],
        ["Supervision compliance (%)", "", "", ""],
        ["Training compliance (%)", "", "", ""],
    ])
    h2(doc, "1.2 Progress and experiences of each child")
    box(doc, "Summary per child: education, health, family time, identity, "
             "enjoyment and achievement, independence, risks", 4.5,
        "Anonymise (use initials) if the report will be shared beyond those "
        "with a need to know.")
    h2(doc, "1.3 Quality Standards self-assessment")
    grid_table(doc, ["Quality Standard (Regs 6–14)", "Strengths",
                     "Areas to improve"],
               [5.4, 6.0, 6.0], blank_rows=0, prefill=[
        ["Quality and purpose of care (Reg 6)", "", ""],
        ["Children's views, wishes and feelings (Reg 7)", "", ""],
        ["Education (Reg 8)", "", ""],
        ["Enjoyment and achievement (Reg 9)", "", ""],
        ["Health and wellbeing (Reg 10)", "", ""],
        ["Positive relationships (Reg 11)", "", ""],
        ["Protection of children (Reg 12)", "", ""],
        ["Leadership and management (Reg 13)", "", ""],
        ["Care planning (Reg 14)", "", ""],
    ])

    h1(doc, "2. Feedback and opinions of children — Reg 45(2)(b)")
    box(doc, "What children have told us this period and how their views "
             "shaped the running of the home", 3.5,
        "House meetings, key work, surveys, complaints, advocate feedback, "
        "day-to-day comments. Give examples of 'you said, we did'.")
    box(doc, "Feedback from parents, placing authorities, IROs, schools and "
             "other professionals", 2.5)

    h1(doc, "3. Complaints — actions taken")
    box(doc, "Complaints received, outcomes and what changed as a result",
        2.5)

    h1(doc, "4. Environment, staffing and leadership")
    box(doc, "Premises, location risk assessment, staffing position and "
             "development of the team", 2.5)

    h1(doc, "5. Judgement and improvement plan")
    box(doc, "Overall evaluation: is the home delivering good quality care "
             "and achieving positive outcomes? What needs to improve?", 2.5)
    grid_table(doc, ["No.", "Action", "Owner", "By when", "Review"],
               [1.2, 8.2, 2.8, 2.6, 2.6], blank_rows=6, row_h=0.7)
    sig_table(doc, ["Registered manager", "Responsible individual"])
    return doc


def t_daily_log_template():
    doc = new_doc("Daily Log")
    guidance(doc, "Daily logs are part of the child's record and may be read "
                  "by the child now or in the future, and by inspectors, "
                  "courts and the Reg 44 visitor. Write to the child, or at "
                  "least for the child.")
    h1(doc, "Recording standards — read before writing")
    bullets(doc, [
        "Be factual: record what you saw and heard, not your "
        "interpretation. Say \"Jay slammed the door and shouted\" not "
        "\"Jay was attention-seeking\".",
        "Separate fact from opinion; if you give a professional opinion, "
        "label it as such and give your reasons.",
        "Use the child's own words in quotation marks for anything "
        "significant they say.",
        "Avoid jargon, labels and abbreviations; be respectful — the child "
        "may read this.",
        "Record positives and achievements, not just problems.",
        "Never leave gaps: record every shift, sign and time every entry, "
        "and never alter an entry after the fact — add a dated correction.",
    ])
    h1(doc, "Daily record")
    field_table(doc, ["Child / young person", "Date",
                      "Staff on shift (early / late / night)"],
                header="Day details", label_w=5.0)
    for label, height, hint in [
        ("Morning (waking, mood, breakfast, getting to school / activity)",
         2.8, ""),
        ("Daytime (school / college / activity — attendance, feedback, "
         "appointments)", 2.8, ""),
        ("Evening (meal, activities, family time / contact, peers, mood)",
         2.8, ""),
        ("Night (settling time, checks as per plan, any waking)", 2.2, ""),
        ("Health, medication and wellbeing notes", 1.8,
         "Cross-reference the MAR sheet for medication detail."),
        ("Positives and achievements today", 1.8, ""),
        ("Concerns, incidents or safeguarding notes", 1.8,
         "Cross-reference incident / intervention records where completed."),
    ]:
        box(doc, label, height, hint)
    grid_table(doc, ["Entry written by", "Role", "Time", "Signature"],
               [6.4, 4.0, 3.0, 4.0], blank_rows=3, row_h=0.7)
    return doc


def t_handover_sheet():
    doc = new_doc("Shift Handover Sheet")
    guidance(doc, "Complete at every shift change. Both the outgoing and "
                  "incoming senior sign to confirm the handover took place "
                  "and key risks were shared.")
    field_table(doc, [
        "Date", "Shift handing over / shift starting",
        "Outgoing staff", "Incoming staff",
    ], header="Shift details", label_w=5.0)
    h1(doc, "Children / young people")
    grid_table(doc, ["Name", "Present / whereabouts",
                     "Key updates this shift (mood, incidents, positives)",
                     "Actions for next shift"],
               [3.0, 3.2, 6.6, 4.6], blank_rows=5, row_h=1.1)
    h1(doc, "Essential checks")
    field_table(doc, [
        ("Medication", "Given as prescribed? Due next shift? Stock issues? "
                       "Cross-reference MAR"),
        "Appointments today / tomorrow (who, where, transport)",
        "Incidents / accidents this shift (records completed?)",
        "Missing / safeguarding updates and current risk changes",
        "Family time / contact arrangements upcoming",
        "Money: petty cash balance, pocket money, receipts",
        "Vehicle: fuel, condition, keys",
        "Maintenance / health and safety issues reported",
        "Messages (calls from social workers, school, family)",
        "Outstanding tasks handed over",
    ], row_h=0.9)
    h1(doc, "Sign-off")
    sig_table(doc, ["Outgoing senior (handed over)",
                    "Incoming senior (accepted)"])
    return doc


def t_friendship_associates_log():
    doc = new_doc("Friendship & Associates Log", landscape=True)
    guidance(doc, "Keep an up-to-date picture of who each child spends time "
                  "with, on and offline. This supports protection from "
                  "exploitation (CSE / CCE, county lines) under Regulation 12 "
                  "and informs missing episode responses. Share relevant "
                  "concerns with the social worker and police as appropriate.")
    field_table(doc, ["Child / young person", "Started (date)",
                      "Reviewed (dates)"], header="Log details", label_w=5.0,
                total_w=LANDSCAPE_W)
    grid_table(doc, ["Date first noted", "Name / nickname",
                     "Approx. age", "How they know each other / where they "
                     "meet (incl. online)", "Address / contact if known",
                     "Concerns? (Y/N — detail)", "Action taken / shared with",
                     "Initials"],
               [2.2, 2.8, 1.6, 4.4, 3.2, 4.8, 5.6, 1.5],
               blank_rows=8, row_h=1.2, font_size=8.5)
    doc.add_paragraph()
    h1(doc, "Prompts for concern")
    bullets(doc, [
        "Significantly older friends or unknown adults; reluctance to say "
        "who they are meeting.",
        "New phones, money, clothes, vapes or gifts without explanation.",
        "Being collected in cars by unknown people; travel to other towns.",
        "Secretive online contacts, new accounts, sending images.",
        "Peers linked to previous missing episodes, drugs or offending.",
    ])
    guidance(doc, "If concerns suggest exploitation, follow your "
                  "safeguarding procedure immediately — do not wait for "
                  "patterns to build in this log.")
    return doc


def t_young_persons_guide_skeleton():
    doc = new_doc("Young Person's Guide — Skeleton")
    guidance(doc, "A skeleton for your children's guide. Rewrite every "
                  "section in language, format and length that suits the "
                  "ages, understanding and communication needs of the "
                  "children you care for — add photos, symbols or an "
                  "easy-read / audio version as needed. Give it to every "
                  "child when they arrive and go through it with them.")
    sections = [
        ("Welcome to [name of home]!", "A warm welcome. Who lives here, "
         "what kind of place it is, and a sentence about what we're here "
         "to do for you."),
        ("The adults who look after you", "Who works here (first names and "
         "photos), what a key worker is, and who the manager is."),
        ("Your bedroom and our home", "About your own room, personalising "
         "it, shared spaces, laundry, and who else lives here."),
        ("How we do things — house routines", "Mealtimes, school days, "
         "bedtimes, phones and internet, having friends round."),
        ("Food", "How meals work, having a say in menus, snacks, special "
         "diets and favourite foods."),
        ("School and learning", "How we support school, homework help, and "
         "what happens if you're not in school right now."),
        ("Pocket money, savings and clothes", "How much, when, savings, and "
         "clothing allowances."),
        ("Seeing your family and friends", "How family time works, phone "
         "calls, and who to talk to if you want it to change."),
        ("Keeping you safe", "How we keep everyone safe, what happens if "
         "you go missing (we will look for you and want you back), and "
         "rules about consequences — what staff can and cannot do."),
        ("Your voice — having your say", "House meetings, your reviews, "
         "your social worker and IRO, and how you can see what's written "
         "about you."),
        ("Not happy? How to complain", "How to complain and who to — the "
         "manager, your social worker, an advocate. Nothing bad will "
         "happen to you for complaining."),
        ("People from outside who check on you", "Your Reg 44 independent "
         "visitor (comes every month and wants to talk to you) and Ofsted "
         "inspectors."),
        ("Important numbers", "Your social worker: [   ]  •  Advocacy "
         "service: [   ]  •  Childline: 0800 1111  •  Ofsted: 0300 123 1231 "
         " •  Children's Commissioner Help at Hand: 0800 528 0731"),
    ]
    for title, hint in sections:
        h1(doc, title)
        box(doc, "Write this section here", 1.8, hint)
    return doc


# -------------------------------------------------------------- CQC builders ---

def t_person_centred_care_plan():
    doc = new_doc("Person-Centred Care Plan")
    guidance(doc, "Written with the person (and those important to them) — "
                  "not about them. Supports Regulation 9 of the Health and "
                  "Social Care Act 2008 (Regulated Activities) Regulations "
                  "2014. Review monthly, after any significant change, and "
                  "involve the person in every review.")
    h1(doc, "1. About me")
    field_table(doc, [
        "Name", "What I like to be called", "Date of birth", "NHS number",
        "GP (name, practice, phone)", "Room / address",
        "Date plan started", "Key worker",
        ("People important to me", "Family, friends, advocate — names, "
                                   "relationships, contact details"),
        ("Who to contact in an emergency", ""),
        ("Legal matters", "LPA (health & welfare / finance), deputy, DNACPR "
                          "/ ReSPECT in place?"),
    ], header="Personal details")
    box(doc, "My life story, background and what matters to me", 3.0,
        "Work, family, faith, culture, achievements, routines that matter, "
        "things I'm proud of.")
    box(doc, "How I communicate and how to communicate with me", 2.2,
        "Language, hearing / sight aids, capacity considerations, how I show "
        "pain or distress.")

    h1(doc, "2. Consent and involvement")
    field_table(doc, [
        ("Does the person consent to this plan of care?", "Yes / No"),
        ("If capacity is in doubt", "Decision-specific MCA assessment "
                                    "completed? Best interests decision "
                                    "recorded? (attach)"),
        "Who was involved in writing this plan",
        ("How the person wants information shared", "and with whom"),
    ], row_h=1.0)

    h1(doc, "3. My needs and outcomes")
    guidance(doc, "Complete each domain: what I can do for myself, what "
                  "support I need, and the outcome we are working towards. "
                  "Cross-reference risk assessments.")
    domains = [
        ("Personal care, dressing and appearance", ""),
        ("Mobility and falls", "Include equipment, transfers, falls history "
                               "and prevention."),
        ("Eating, drinking and nutrition", "Include weight monitoring, MUST "
                                           "score, texture (IDDSI), likes "
                                           "and dislikes."),
        ("Continence", ""),
        ("Skin integrity", "Include pressure risk (e.g. Waterlow), "
                           "repositioning, equipment."),
        ("Medication", "Level of support, self-administration assessment, "
                       "PRN protocols, covert medication (MCA) if any."),
        ("Health conditions and monitoring", "Diagnoses, appointments, what "
                                             "a deterioration looks like and "
                                             "what to do."),
        ("Mental health, memory and emotional wellbeing", ""),
        ("Sleep and night-time support", ""),
        ("Social life, activities and things I enjoy", ""),
        ("Religion, culture and relationships", ""),
        ("End of life wishes", "Only where the person wishes to discuss; "
                               "advance care plan, funeral wishes, who to "
                               "involve."),
    ]
    for name, hint in domains:
        h2(doc, name)
        field_table(doc, [
            ("What I can do for myself", hint),
            "Support I need and how I want it given",
            "Outcome / goal we are working towards",
            "Linked risk assessment(s)",
        ], row_h=0.95, label_w=5.5)

    h1(doc, "4. Reviews")
    grid_table(doc, ["Date", "Reviewed with (person / family / advocate)",
                     "Changes made", "Next review", "Signature"],
               [2.2, 4.6, 5.6, 2.4, 2.6], blank_rows=6, row_h=0.8)
    sig_table(doc, ["The person (or representative)", "Key worker",
                    "Manager"])
    return doc


def t_one_page_profile():
    doc = new_doc("One-Page Profile")
    guidance(doc, "A single page that tells anyone supporting this person "
                  "what matters most. Keep it current, keep it positive, and "
                  "write it in the person's own words wherever possible. "
                  "Display or file it where staff will actually see it.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    fixed_layout(table)
    cell_text(table.rows[0].cells[0], "Name and what I like to be called:",
              bold=True, fill=LIGHT_FILL)
    cell_text(table.rows[0].cells[1], "Photo (with my consent)", italic=True,
              size=8.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    table.rows[0].cells[0].width = Cm(12.4)
    table.rows[0].cells[1].width = Cm(5.0)
    row_height(table.rows[0], 3.2)
    doc.add_paragraph()
    box(doc, "What people appreciate about me", 4.0,
        "My qualities, talents and character — ask the people who know and "
        "love me. e.g. 'great sense of humour', 'wonderful grandad', "
        "'always notices when someone is down'.")
    box(doc, "What is important to me", 5.0,
        "The people, routines, possessions, faith, food, activities and "
        "little things that make my day — specific enough that a stranger "
        "could get it right. e.g. 'tea in my own blue mug, milk in first'.")
    box(doc, "How best to support me", 5.0,
        "What good support looks and sounds like, what to avoid, how I "
        "communicate, and how you'll know if I'm unhappy or in pain.")
    field_table(doc, ["Written with", "Date", "Review date"], label_w=4.5)
    return doc


def t_mca_capacity_assessment():
    doc = new_doc("Mental Capacity Assessment (MCA 2005)")
    guidance(doc, "A capacity assessment is decision-specific and "
                  "time-specific. Start from the presumption of capacity "
                  "(principle 1), take all practicable steps to help the "
                  "person decide (principle 2), and remember an unwise "
                  "decision does not mean a lack of capacity (principle 3).")
    h1(doc, "1. Details")
    field_table(doc, [
        "Person's name", "Date of birth", "Address / room",
        "Assessor (name, role)", "Date(s) of assessment",
        "Others present / consulted",
    ], header="Assessment details")
    field_table(doc, [
        ("The specific decision to be made",
         "Be precise — e.g. 'whether to accept support with washing each "
         "morning', not 'care decisions'"),
        ("Why is capacity in doubt for this decision?",
         "What has been observed or said that triggered this assessment?"),
        ("Why now?", "Why does this decision need to be made at this time? "
                     "Could it wait until capacity may be regained?"),
    ], row_h=1.2)

    h1(doc, "2. Supporting the person to decide (principle 2)")
    box(doc, "All practicable steps taken to help the person make this "
             "decision themselves", 2.5,
        "e.g. simple language, pictures, interpreter, hearing aids in and "
        "working, best time of day, familiar people present, quiet setting, "
        "trying on more than one occasion.")

    h1(doc, "3. Stage 1 — the diagnostic test")
    field_table(doc, [
        ("Is there an impairment of, or disturbance in the functioning of, "
         "the mind or brain?", "Yes / No — e.g. dementia, learning "
         "disability, brain injury, delirium, intoxication. It may be "
         "temporary or permanent."),
        ("Evidence", "Diagnosis, observations, information from GP or "
                     "records"),
    ], row_h=1.2)
    guidance(doc, "If No — the person does not lack capacity under the MCA. "
                  "Stop here and record the outcome in section 5.")

    h1(doc, "4. Stage 2 — the functional test")
    guidance(doc, "The person is unable to make the decision if they cannot "
                  "do one or more of the following, because of the "
                  "impairment or disturbance. Record the questions asked and "
                  "the person's actual responses.")
    for label, hint in [
        ("(a) Understand the information relevant to the decision",
         "Include the salient details: what the decision is, why it is "
         "needed, and the likely consequences of deciding either way or "
         "not deciding. Can they understand it? Yes / No — evidence."),
        ("(b) Retain that information",
         "Long enough to make the decision — retention for a short period "
         "only does not by itself mean incapacity. Yes / No — evidence."),
        ("(c) Use or weigh that information as part of making the decision",
         "Can they weigh the pros, cons and consequences, and apply them "
         "to their own situation? Yes / No — evidence."),
        ("(d) Communicate the decision",
         "By any means — talking, writing, sign, blinking. Yes / No — "
         "evidence."),
    ]:
        box(doc, label, 2.6, hint)

    h1(doc, "5. Conclusion")
    field_table(doc, [
        ("Outcome", "On the balance of probabilities, does the person have "
                    "capacity to make THIS decision at THIS time? "
                    "Has capacity / Lacks capacity"),
        ("Is the inability caused by the impairment / disturbance?",
         "The 'causative nexus' — Yes / No"),
        ("Is capacity likely to be regained?",
         "If yes, can the decision wait? When should this be reassessed?"),
        ("If the person lacks capacity",
         "Proceed to a best interests decision (s.4 MCA) — record on the "
         "Best Interests Decision Record. Consider IMCA referral if "
         "unbefriended, and whether the care proposed amounts to a "
         "deprivation of liberty."),
        "Review date for this assessment",
    ], row_h=1.2)
    sig_table(doc, ["Assessor", "Manager (countersign)"])
    return doc


def t_best_interests_decision_record():
    doc = new_doc("Best Interests Decision Record (MCA 2005, s.4)")
    guidance(doc, "Complete only after a decision-specific capacity "
                  "assessment has concluded the person lacks capacity for "
                  "this decision. The decision-maker must reasonably believe "
                  "the act or decision is in the person's best interests.")
    h1(doc, "1. Details")
    field_table(doc, [
        "Person's name", "Date of birth",
        "Decision to be made (specific)",
        "Date and outcome of capacity assessment (attach)",
        ("Decision-maker", "Name and role — the person proposing the act, "
                           "e.g. carer for care acts, GP for treatment; an "
                           "attorney / deputy if the decision is within "
                           "their authority"),
    ], header="Decision details", row_h=1.0)

    h1(doc, "2. Section 4 checklist")
    grid_table(doc, ["Consideration", "How it was considered / evidence"],
               [7.4, 10.0], blank_rows=0, prefill=[
        ["No assumptions made merely on the basis of age, appearance, "
         "condition or behaviour", ""],
        ["All relevant circumstances considered", ""],
        ["Could the decision be delayed until the person regains capacity?",
         ""],
        ["The person has been encouraged and enabled to participate as "
         "fully as possible", ""],
        ["The person's past and present wishes and feelings (including any "
         "written statements made when they had capacity)", ""],
        ["The person's beliefs and values likely to influence the decision",
         ""],
        ["Other factors the person would be likely to consider", ""],
        ["If the decision concerns life-sustaining treatment: not motivated "
         "by a desire to bring about death", ""],
    ])
    for r in doc.tables[-1].rows[1:]:
        row_height(r, 1.1)

    h1(doc, "3. Consultation")
    grid_table(doc, ["Consulted (name)", "Relationship / role",
                     "Their views on the person's best interests", "Date"],
               [3.4, 3.2, 8.6, 2.2], blank_rows=5, row_h=1.0)
    field_table(doc, [
        ("Attorney (LPA) or Court-appointed deputy?",
         "If yes — do they have authority for this decision?"),
        ("IMCA instructed?", "Required for serious medical treatment or "
                             "long-term accommodation moves where the person "
                             "has no one appropriate to consult"),
    ], row_h=1.0)

    h1(doc, "4. Options considered")
    guidance(doc, "Consider whether the purpose can be achieved in a way "
                  "that is less restrictive of the person's rights and "
                  "freedom of action (principle 5).")
    grid_table(doc, ["Option (include 'do nothing')", "Benefits", "Risks / "
                     "burdens", "How restrictive?"],
               [4.6, 4.4, 4.4, 4.0], blank_rows=4, row_h=1.2)

    h1(doc, "5. The decision")
    box(doc, "Decision made and full rationale — why this option is in the "
             "person's best interests", 3.0)
    field_table(doc, [
        "Any disagreement and how it will be resolved "
        "(mediation, second opinion, Court of Protection)",
        ("Does the plan amount to a deprivation of liberty?",
         "If yes — DoLS / Court of Protection authorisation required"),
        "How and when this decision will be reviewed",
        "Decision date",
    ], row_h=1.0)
    sig_table(doc, ["Decision-maker", "Manager (countersign)"])
    return doc


def t_mar_audit_tool():
    doc = new_doc("MAR Chart Audit Tool")
    guidance(doc, "Audit a sample of Medication Administration Records at "
                  "least monthly. For each question tick Yes, No or N/A; "
                  "every 'No' requires an action in the plan at the end. "
                  "Supports Regulation 12(2)(g) — the safe management of "
                  "medicines.")
    field_table(doc, [
        "Service / unit", "Auditor (name, role)", "Date of audit",
        ("Sample", "Number of MAR charts reviewed / people's initials"),
        "Period covered",
    ], header="Audit details", label_w=5.0)

    h1(doc, "Audit questions")
    questions = [
        "Each MAR has the person's name, photo, date of birth and GP "
        "recorded",
        "Allergies and sensitivities are recorded on every MAR (or 'none "
        "known')",
        "Every administration is signed at the time of administration",
        "There are no unexplained gaps in administration signatures",
        "Omission codes are used correctly and reasons for refusal recorded",
        "Handwritten / mid-cycle entries are signed and witnessed by a "
        "second member of staff",
        "Quantities received, carried over and returned are recorded and "
        "reconcile with stock",
        "Running stock balances of boxed medication are correct (sample "
        "count)",
        "PRN medicines have an up-to-date person-specific protocol (dose, "
        "indication, interval, max in 24h)",
        "PRN administrations record the reason given and the outcome / "
        "effectiveness",
        "Variable dose administrations record the actual dose given",
        "Time-sensitive medicines (e.g. Parkinson's) are given at the "
        "prescribed times",
        "Controlled drugs register entries are double-signed and the "
        "balance is correct (sample count)",
        "Medication room / trolley temperature recorded daily and below "
        "25°C",
        "Fridge temperature recorded daily and within 2–8°C, with action "
        "taken when out of range",
        "Medicines are in date, labelled, and opening dates recorded on "
        "limited-life items (eye drops, insulin, liquids)",
        "Covert administration (if any) is supported by an MCA best "
        "interests decision and pharmacist advice",
        "Self-administration (if any) is supported by a risk assessment",
        "Staff administering have completed medication training and "
        "competency assessment in date",
        "Medication errors since last audit were recorded, reported and "
        "actioned",
    ]
    grid_table(doc, ["#", "Audit question", "Yes", "No", "N/A",
                     "Comments / action required"],
               [0.9, 8.3, 1.1, 1.1, 1.1, 4.9], blank_rows=0,
               prefill=[[str(i + 1), q, "", "", "", ""]
                        for i, q in enumerate(questions)], font_size=9)

    h1(doc, "Action plan")
    grid_table(doc, ["No.", "Action", "Owner", "By when", "Completed"],
               [1.2, 8.6, 2.8, 2.4, 2.4], blank_rows=5, row_h=0.7)
    sig_table(doc, ["Auditor", "Registered manager"])
    return doc


def t_medication_error_record():
    doc = new_doc("Medication Error Record")
    guidance(doc, "Complete as soon as the error is discovered — the "
                  "priority is always the person's safety first, paperwork "
                  "second. Never conceal an error; a no-blame reporting "
                  "culture protects people.")
    h1(doc, "1. The error")
    field_table(doc, [
        "Person affected", "Date of birth",
        "Date and time error occurred (if known)",
        "Date and time error discovered, and by whom",
        ("Medication(s) involved", "Name, strength, dose, route"),
        ("Type of error", "Missed dose / wrong dose / wrong medicine / "
                          "wrong person / wrong time / wrong route / "
                          "unauthorised administration / recording error / "
                          "stock or storage error / other"),
    ], header="Error details", row_h=1.0)
    box(doc, "What happened — factual account", 3.0)

    h1(doc, "2. Immediate actions — person's safety")
    field_table(doc, [
        "Person checked — observations and current condition",
        ("Medical advice sought", "GP / NHS 111 / pharmacist / 999 — who "
                                  "was contacted, when, by whom"),
        ("Advice given", "Record verbatim, including any monitoring "
                         "instructions and red flags to watch for"),
        "Monitoring carried out (what, how often, findings)",
        "Person / family informed (who, when — see also duty of candour "
        "if harm occurred)",
    ], row_h=1.1)

    h1(doc, "3. Notifications")
    check_table(doc, [
        "Manager on call / registered manager informed",
        "GP informed",
        "CQC statutory notification submitted (if harm / abuse threshold "
        "met)",
        "Local authority safeguarding referral (if threshold met)",
        "Commissioners / funding authority informed (if required)",
        "Duty of candour process started (if notifiable safety incident)",
    ], first_col="Notification", extra=("Date / time", "By whom", "Ref"),
        extra_w=2.8)

    h1(doc, "4. Staff involved")
    field_table(doc, [
        "Staff member(s) involved in the error",
        "Medication training and competency in date? (dates)",
        ("Immediate practice action", "e.g. supervised practice, retraining, "
                                      "competency reassessment before next "
                                      "medication round"),
    ], row_h=1.0)

    h1(doc, "5. Reflection and learning")
    box(doc, "Root cause and contributory factors", 2.5,
        "e.g. interruptions during the round, look-alike packaging, unclear "
        "MAR, staffing, process gaps — focus on systems, not blame.")
    grid_table(doc, ["Action to prevent recurrence", "Owner", "By when",
                     "Completed"],
               [9.4, 3.0, 2.4, 2.6], blank_rows=4, row_h=0.7)
    sig_table(doc, ["Staff member reporting", "Registered manager (review)"])
    return doc


def t_incident_accident_form():
    doc = new_doc("Incident / Accident Form")
    guidance(doc, "Use for accidents, incidents and near misses involving "
                  "the people we support, staff or visitors. Complete before "
                  "the end of the shift.")
    h1(doc, "1. Details")
    field_table(doc, [
        ("Person(s) involved", "Name(s) and whether resident / staff / "
                               "visitor / other"),
        "Date and time of incident", "Exact location",
        ("Type", "Fall / injury / near miss / behaviour incident / "
                 "medication (also complete medication error record) / "
                 "absconding / property damage / security / other"),
        "Witnesses (names and roles)",
    ], header="Incident details", row_h=1.0)
    box(doc, "What happened — factual description", 3.5,
        "What was seen and heard, what the person was doing beforehand, "
        "sequence of events. Attach witness statements if taken.")
    box(doc, "Injuries sustained and treatment given", 2.2,
        "Mark on a body map where applicable. Record 'no visible injury' "
        "checks and any later observations (e.g. head injury monitoring).")
    field_table(doc, [
        "First aid given by / emergency services called (time, outcome)",
        "GP / hospital attendance (details)",
    ], row_h=0.9)

    h1(doc, "2. Immediate actions taken")
    box(doc, "Actions to make the situation safe and support those involved",
        2.2)

    h1(doc, "3. RIDDOR consideration")
    guidance(doc, "The Reporting of Injuries, Diseases and Dangerous "
                  "Occurrences Regulations 2013 apply to work-related "
                  "incidents. Consider with the manager.")
    field_table(doc, [
        ("Is this RIDDOR reportable?", "Death / specified injury / over-7-"
                                       "day incapacitation of a worker / "
                                       "injury taking a non-worker to "
                                       "hospital / occupational disease / "
                                       "dangerous occurrence"),
        "If yes — reported to HSE (date, method, reference)",
    ], row_h=1.1)

    h1(doc, "4. Notifications")
    check_table(doc, [
        "Registered manager informed",
        "Family / representative informed",
        "GP / healthcare professional informed",
        "CQC notification (Reg 18, CQC (Registration) Regs 2009) if "
        "threshold met",
        "Safeguarding referral if abuse or neglect suspected",
        "Ofsted notification (children's services, Reg 40) if applicable",
    ], first_col="Notification", extra=("Date", "By whom"), extra_w=3.0)

    h1(doc, "5. Manager review and learning")
    field_table(doc, [
        "Risk assessment(s) / care plan reviewed and updated? (which, date)",
        "Falls / incident pattern check completed?",
        "Learning and actions arising",
    ], row_h=1.1)
    sig_table(doc, ["Staff member completing form",
                    "Registered manager (review)"])
    return doc


def t_complaints_log():
    doc = new_doc("Complaints Log", landscape=True)
    guidance(doc, "A running record of all complaints, including verbal "
                  "complaints and grumbles, supporting Regulation 16. Keep "
                  "individual complaint files for detail; use this log for "
                  "oversight, trends and evidence of learning.")
    field_table(doc, ["Service", "Year", "Log maintained by"],
                header="Log details", label_w=5.0, total_w=LANDSCAPE_W)
    h1(doc, "Log")
    grid_table(doc, ["Ref", "Date received", "Complainant (and on behalf "
                     "of)", "How received", "Summary of complaint",
                     "Acknowledged (date)", "Investigated by",
                     "Outcome (upheld / partly / not upheld)",
                     "Response sent", "Actions / learning", "Closed"],
               [1.2, 1.9, 2.6, 1.7, 4.1, 1.9, 2.1, 2.3, 1.8, 4.3, 2.2],
               blank_rows=8, row_h=1.3, font_size=8)
    h1(doc, "Timescales — our standards")
    bullets(doc, [
        "Acknowledge every complaint within 3 working days.",
        "Investigate and respond in full within 28 days, or write "
        "explaining the delay and the new date.",
        "Always include in the response: what we found, what we will "
        "change, and how to escalate — the Local Government and Social "
        "Care Ombudsman (0300 061 0614) for adult social care; complainants "
        "may also tell CQC at any time (03000 616161).",
        "Review this log quarterly for themes; feed learning into team "
        "meetings and the quality audit cycle.",
    ])
    return doc


def t_duty_of_candour_record():
    doc = new_doc("Duty of Candour Record (Regulation 20)")
    guidance(doc, "Use when a notifiable safety incident has, or may have, "
                  "occurred: an unintended or unexpected incident that in "
                  "the reasonable opinion of a healthcare professional "
                  "resulted in (or could result in) death, severe harm, "
                  "moderate harm or prolonged psychological harm. Being "
                  "open and saying sorry is not an admission of liability.")
    h1(doc, "1. The incident")
    field_table(doc, [
        "Person affected", "Date of birth", "Date and time of incident",
        "Incident / accident record reference",
        ("Degree of harm", "Death / severe / moderate / prolonged "
                           "psychological harm — and clinical opinion "
                           "source"),
        ("Relevant person", "The person themselves, or their representative "
                            "if the person is deceased or lacks capacity"),
    ], header="Notifiable safety incident details", row_h=1.0)
    box(doc, "What happened — all the facts known at this stage", 2.8)

    h1(doc, "2. Verbal notification (as soon as reasonably practicable)")
    field_table(doc, [
        "Date, time and place of verbal notification",
        "Given by (name, role)",
        ("Account given", "A true account of all facts known so far"),
        ("Apology given", "Record that a sincere apology was expressed"),
        "What we said would happen next (further enquiries / "
        "investigation)",
        "Questions asked by the relevant person and answers given",
        "Support offered (emotional and practical, advocacy, interpreter)",
        ("If notification was not practicable", "Record why, and attempts "
                                                "made"),
    ], row_h=1.0)

    h1(doc, "3. Written follow-up")
    field_table(doc, [
        "Date written notification sent (attach a copy)",
        ("Contents checklist", "Account of the incident; apology; details "
                               "of further enquiries; how results will be "
                               "shared"),
        "Results of further enquiries / investigation shared (date, how)",
        "Any further correspondence (dates)",
    ], row_h=1.0)

    h1(doc, "4. Investigation and learning")
    box(doc, "Findings of the investigation and actions taken to prevent "
             "recurrence", 2.8)
    field_table(doc, [
        "CQC statutory notification submitted (date, reference)",
        "Learning shared with the team (how, date)",
    ], row_h=0.9)
    h1(doc, "5. Sign-off")
    sig_table(doc, ["Staff member leading candour process",
                    "Registered manager"])
    guidance(doc, "Keep this record with copies of all correspondence — CQC "
                  "will expect to see a documented audit trail of "
                  "compliance with each stage of Regulation 20.")
    return doc


def t_statement_of_purpose_skeleton():
    doc = new_doc("Statement of Purpose — Skeleton")
    guidance(doc, "Every registered provider must have a statement of "
                  "purpose containing the matters in Schedule 3 of the Care "
                  "Quality Commission (Registration) Regulations 2009 (Reg "
                  "12). Keep it under review, and notify CQC of any changes "
                  "within 28 days.")
    sections = [
        ("1. Provider details", "Full name of the registered provider, "
         "legal status (individual / partnership / organisation), business "
         "address, telephone and email, and CQC provider ID."),
        ("2. Aims and objectives", "What the service sets out to achieve "
         "for the people it supports, its ethos and values, and how it "
         "delivers care in line with the fundamental standards."),
        ("3. Regulated activities and service types", "The regulated "
         "activities carried on (e.g. accommodation for persons who "
         "require nursing or personal care; personal care), the service "
         "types (e.g. care home without nursing, domiciliary care agency), "
         "and any conditions of registration."),
        ("4. The needs we meet", "Service user bands (e.g. older people, "
         "dementia, learning disabilities or autistic people, mental "
         "health, physical disability), age ranges, and any specialisms. "
         "State the number of people the service is registered for."),
        ("5. Locations", "For each location: name, address, CQC location "
         "ID, description of the premises and facilities, communal and "
         "accessible spaces, and the regulated activities carried on "
         "there."),
        ("6. Staffing and management structure", "Overall structure from "
         "nominated individual to care staff, staffing levels by day and "
         "night, qualifications and training approach, and use of bank or "
         "agency staff."),
        ("7. The registered manager", "Name, contact details, relevant "
         "qualifications and experience, and the locations they manage."),
        ("8. How care is planned and reviewed", "Assessment before "
         "admission, person-centred planning, involvement of people and "
         "families, and review cycles."),
        ("9. Concerns, complaints and compliments", "How to complain, "
         "timescales, and escalation to the Local Government and Social "
         "Care Ombudsman; how feedback shapes the service."),
        ("10. Contact details", "Service contact details and CQC's "
         "details: Care Quality Commission, Citygate, Gallowgate, "
         "Newcastle upon Tyne NE1 4PA — 03000 616161 — www.cqc.org.uk."),
    ]
    for title, hint in sections:
        h1(doc, title)
        box(doc, "Write this section here", 2.0, hint)
    field_table(doc, ["Date of this version", "Approved by",
                      "Next review due"], label_w=5.0)
    return doc


def t_service_user_guide_skeleton():
    doc = new_doc("Service User Guide — Skeleton")
    guidance(doc, "A plain-English welcome guide for people using the "
                  "service and their families. Produce accessible versions "
                  "(large print, easy read, other languages) to meet the "
                  "Accessible Information Standard. Give a copy to everyone "
                  "before or on the day their care starts.")
    sections = [
        ("Welcome", "A short, warm introduction: who we are, our values, "
         "and what people can expect from us."),
        ("About our service", "What we provide, who we support, and our "
         "CQC registration (including our latest rating and where to read "
         "the report)."),
        ("Our team", "Who's who — manager, senior staff, key workers — and "
         "how staff are trained, supervised and DBS checked."),
        ("Your care and support plan", "How we assess needs, how the plan "
         "is agreed with you, how often it is reviewed, and how you and "
         "your family are involved."),
        ("Daily life", "Meals and choices, activities, visitors (no "
         "restrictions without good reason), going out, pets, phone and "
         "internet."),
        ("Your room and our building / Our visits to your home",
         "Adapt to setting: personalising rooms, laundry, keys; or for "
         "home care — visit times, what happens if a carer is late, keys "
         "and access arrangements."),
        ("Fees and charges", "What is included, what costs extra, invoicing "
         "arrangements, and what happens if funding changes. Reference the "
         "contract / terms."),
        ("Keeping you safe", "Safeguarding — what we do if we are worried "
         "someone is being harmed, and who you can talk to. Include the "
         "local authority safeguarding number."),
        ("Your medicines", "How we support with medicines, your right to "
         "manage your own where safe, and how errors are handled openly."),
        ("Your rights", "Dignity, privacy, choice and independence; "
         "consent and the Mental Capacity Act; advocacy services and how "
         "to contact them; how to see your records (data protection)."),
        ("Compliments, comments and complaints", "How to complain, our "
         "timescales, and who else can help: the Local Government and "
         "Social Care Ombudsman (0300 061 0614) and CQC (03000 616161). "
         "Complaining will never affect your care."),
        ("Useful contacts", "Service phone / out-of-hours, local authority, "
         "advocacy, Healthwatch, CQC."),
    ]
    for title, hint in sections:
        h1(doc, title)
        box(doc, "Write this section here", 1.8, hint)
    return doc


# ------------------------------------------------------------ shared builders ---

def t_supervision_record_staff():
    doc = new_doc("Supervision Record — Staff")
    guidance(doc, "One-to-one supervision should be regular (at least "
                  "every 4–6 weeks for care staff, or per your policy), "
                  "planned and private. It is a two-way, confidential "
                  "conversation — but confidentiality never extends to "
                  "safeguarding concerns.")
    field_table(doc, [
        "Staff member", "Role", "Supervisor", "Date", "Location",
        "Date of last supervision",
    ], header="Supervision details", label_w=5.5)
    h1(doc, "1. Review of previous actions")
    grid_table(doc, ["Action from last supervision", "Update", "Complete?"],
               [7.0, 8.0, 2.4], blank_rows=3, row_h=0.8)
    h1(doc, "2. Wellbeing")
    box(doc, "How are you? Workload, work-life balance, health, anything "
             "affecting you at work", 2.5,
        "Include lone working, night working and support needs.")
    h1(doc, "3. Practice — the people we support")
    box(doc, "Caseload / key children or key residents: progress, "
             "concerns, plans and recording up to date?", 3.5,
        "Discuss each allocated person briefly; record decisions and "
        "rationale.")
    h1(doc, "4. Safeguarding")
    box(doc, "Any safeguarding concerns, incidents, allegations or "
             "whistleblowing matters (reminder of how to raise concerns "
             "given)", 2.2)
    h1(doc, "5. Performance and feedback")
    box(doc, "What's going well; feedback both ways; any performance "
             "concerns and support agreed", 2.5)
    h1(doc, "6. Training and development")
    box(doc, "Training completed / due, qualification progress, career "
             "development", 2.0)
    h1(doc, "7. Agreed actions")
    grid_table(doc, ["Action", "Who", "By when"],
               [10.4, 3.4, 3.6], blank_rows=4, row_h=0.7)
    field_table(doc, ["Date of next supervision"], label_w=5.5)
    sig_table(doc, ["Staff member", "Supervisor"])
    guidance(doc, "Both parties should agree the record is accurate; the "
                  "staff member should receive a copy.")
    return doc


def t_supervision_record_manager():
    doc = new_doc("Supervision Record — Manager")
    guidance(doc, "For supervision of the registered manager by the "
                  "responsible individual / nominated individual or senior "
                  "leader. Focus on the manager's oversight of the service "
                  "as well as their own support and development.")
    field_table(doc, [
        "Manager", "Service", "Supervisor (RI / NI / senior leader)",
        "Date", "Date of last supervision",
    ], header="Supervision details", label_w=5.5)
    h1(doc, "1. Review of previous actions")
    grid_table(doc, ["Action from last supervision", "Update", "Complete?"],
               [7.0, 8.0, 2.4], blank_rows=3, row_h=0.8)
    h1(doc, "2. Manager's wellbeing and workload")
    box(doc, "How are you? On-call load, hours, resilience, support needs",
        2.2)
    h1(doc, "3. Safeguarding and incident oversight")
    box(doc, "Open safeguarding matters, allegations (LADO referrals), "
             "serious incidents, notifications made to the regulator, "
             "patterns and lessons", 2.5)
    h1(doc, "4. Quality and compliance")
    box(doc, "Audit findings, action plans, inspection readiness "
             "(Ofsted / CQC), complaints and compliments, feedback from "
             "people supported", 2.5)
    h1(doc, "5. Staffing")
    box(doc, "Vacancies and recruitment, sickness, agency use, supervision "
             "and appraisal compliance, training compliance, team morale, "
             "any performance / disciplinary matters", 2.5)
    h1(doc, "6. Resources and environment")
    box(doc, "Budget position, occupancy / referrals, premises and "
             "maintenance, health and safety", 2.2)
    h1(doc, "7. Manager's development")
    box(doc, "Leadership development, qualifications (e.g. Level 5), peer "
             "support, career discussion", 2.0)
    h1(doc, "8. Agreed actions")
    grid_table(doc, ["Action", "Who", "By when"],
               [10.4, 3.4, 3.6], blank_rows=4, row_h=0.7)
    field_table(doc, ["Date of next supervision"], label_w=5.5)
    sig_table(doc, ["Manager", "Supervisor"])
    return doc


def t_team_meeting_agenda_minutes():
    doc = new_doc("Team Meeting — Agenda & Minutes")
    guidance(doc, "One document for both agenda and minutes. Circulate the "
                  "agenda in advance; record decisions and actions, not a "
                  "transcript. Staff who were absent must read and sign.")
    field_table(doc, [
        "Service / team", "Date and time", "Chair", "Minute taker",
        "Present", "Apologies",
    ], header="Meeting details", label_w=4.5)
    h1(doc, "Agenda and minutes")
    items = [
        "1. Welcome, confidentiality reminder and apologies",
        "2. Minutes and actions from last meeting",
        "3. Safeguarding — concerns, incidents, notifications, learning",
        "4. The people we support — updates, plans and risk changes "
        "(per person)",
        "5. Health and safety, premises and maintenance",
        "6. Medication — errors, audits, changes",
        "7. Policies and procedure updates for briefing",
        "8. Quality — audit findings, complaints, compliments, feedback",
        "9. Rotas, leave and staffing",
        "10. Training and development",
        "11. Staff wellbeing and recognition",
        "12. Any other business",
    ]
    grid_table(doc, ["Agenda item", "Discussion and decisions"],
               [6.2, 11.2], blank_rows=0, prefill=[[i, ""] for i in items])
    for r in doc.tables[-1].rows[1:]:
        row_height(r, 1.5)
    h1(doc, "Action log")
    grid_table(doc, ["No.", "Action", "Owner", "By when", "Status"],
               [1.2, 8.8, 2.8, 2.4, 2.2], blank_rows=6, row_h=0.7)
    field_table(doc, ["Date, time and venue of next meeting"], label_w=6.5)
    h1(doc, "Read and understood (staff not present)")
    grid_table(doc, ["Name", "Signature", "Date"],
               [6.8, 6.8, 3.8], blank_rows=4, row_h=0.7)
    return doc


def t_reflective_practice_log():
    doc = new_doc("Reflective Practice Log (Gibbs' Cycle)")
    guidance(doc, "Use after a significant piece of work, incident or "
                  "interaction — positive or difficult. Reflection is "
                  "evidence for supervision, appraisal, the Care Certificate "
                  "and professional registration / CPD.")
    field_table(doc, [
        "Name", "Role", "Date of reflection",
        ("Situation reflected on", "Brief title, e.g. 'supporting R after "
                                   "the missing episode on 12/07'"),
    ], header="Details", label_w=5.5)
    stages = [
        ("1. Description — what happened?", 2.6,
         "Just the facts: who was involved, what happened, in what order? "
         "No conclusions yet."),
        ("2. Feelings — what were you thinking and feeling?", 2.4,
         "Before, during and after. What do you think others were "
         "feeling?"),
        ("3. Evaluation — what was good and bad about the experience?", 2.4,
         "What worked? What didn't? What did you and others contribute, "
         "positively or negatively?"),
        ("4. Analysis — what sense can you make of it?", 2.6,
         "Why did things happen the way they did? Link to training, "
         "policy, research or approaches (e.g. trauma-informed practice, "
         "person-centred values)."),
        ("5. Conclusion — what else could you have done?", 2.2,
         "What have you learned about yourself, your practice, the "
         "person?"),
        ("6. Action plan — what will you do differently next time?", 2.2,
         "Concrete, specific commitments. What support or training do you "
         "need?"),
    ]
    for label, height, hint in stages:
        box(doc, label, height, hint)
    field_table(doc, [
        "Discussed in supervision on (date)",
        "Added to development plan / CPD record?",
    ], row_h=0.8)
    sig_table(doc, ["Staff member", "Supervisor (if discussed)"])
    return doc


def t_induction_checklist_care_certificate():
    doc = new_doc("Induction Checklist — Care Certificate")
    guidance(doc, "Structured induction for new care staff, mapped to the "
                  "15 standards of the Care Certificate. The Care "
                  "Certificate should normally be completed within 12 weeks. "
                  "New staff must not work unsupervised until the manager is "
                  "satisfied they are safe to do so.")
    field_table(doc, [
        "New staff member", "Role", "Start date", "Inductor / mentor",
        "Care Certificate target completion date (12 weeks)",
    ], header="Induction details", label_w=6.5)

    h1(doc, "Part 1 — First week essentials")
    check_table(doc, [
        "Tour of the service; introductions to people supported and team",
        "Contract, ID badge, IT accounts and rota explained",
        "DBS and references confirmed cleared by manager before "
        "unsupervised work",
        "Fire safety: evacuation procedure, assembly point, PEEPs, "
        "extinguisher points",
        "Emergency procedures: first aid arrangements, on-call system",
        "Key policies read and signed: safeguarding, whistleblowing, code "
        "of conduct, medication, health & safety, data protection, social "
        "media, lone working",
        "Safeguarding contacts displayed / provided (DSL, local authority, "
        "LADO where applicable)",
        "Infection prevention: hand hygiene and PPE use demonstrated",
        "Moving and handling: no assisting until trained",
        "Introduction to care plans and recording systems",
        "Supervision and probation arrangements explained",
    ], extra=("Date", "Initials"))

    h1(doc, "Part 2 — The 15 Care Certificate standards")
    standards = [
        "1. Understand your role",
        "2. Your personal development",
        "3. Duty of care",
        "4. Equality and diversity",
        "5. Work in a person centred way",
        "6. Communication",
        "7. Privacy and dignity",
        "8. Fluids and nutrition",
        "9. Awareness of mental health, dementia and learning disability",
        "10. Safeguarding adults",
        "11. Safeguarding children",
        "12. Basic life support",
        "13. Health and safety",
        "14. Handling information",
        "15. Infection prevention and control",
    ]
    grid_table(doc, ["Care Certificate standard",
                     "Learning / evidence completed (workbook, observation, "
                     "discussion)", "Date achieved", "Assessor initials"],
               [6.0, 6.0, 2.6, 2.8], blank_rows=0,
               prefill=[[s, "", "", ""] for s in standards])
    for r in doc.tables[-1].rows[1:]:
        row_height(r, 0.8)

    h1(doc, "Sign-off")
    field_table(doc, [
        "All standards assessed as met in the workplace (date)",
        "Probation review dates (weeks 4 / 8 / 12)",
    ], row_h=0.8)
    sig_table(doc, ["New staff member", "Assessor / mentor",
                    "Registered manager"])
    return doc


def t_safer_recruitment_checklist():
    doc = new_doc("Safer Recruitment Checklist")
    guidance(doc, "Track each stage for every appointment. Fitness of "
                  "workers requirements: Schedule 2, Children's Homes "
                  "(England) Regulations 2015; Regulation 19 and Schedule 3, "
                  "Health and Social Care Act 2008 (Regulated Activities) "
                  "Regulations 2014. File this checklist with the evidence "
                  "in the staff record.")
    field_table(doc, [
        "Candidate name", "Post applied for", "Recruiting manager",
        "Interview date", "Proposed start date",
    ], header="Appointment details", label_w=6.0)
    h1(doc, "Recruitment stages")
    check_table(doc, [
        "Advert and job pack include safeguarding commitment and state an "
        "enhanced DBS check is required",
        "Full application form received (a CV alone is not accepted)",
        "Full employment history obtained; ALL gaps identified, explored "
        "at interview and explanations recorded",
        "Shortlisting completed by at least two people",
        "Interview included safeguarding / value-based questions; "
        "conducted by a panel including someone trained in safer "
        "recruitment",
        "Identity verified (original documents seen and copied)",
        "Right to work in the UK verified",
        "Enhanced DBS with barred list check received and clear "
        "(certificate number and date recorded) — or risk assessment by "
        "manager where information disclosed",
        "DBS Update Service status check (where applicable)",
        "Two written references obtained directly from referees — one "
        "from the current / most recent employer — and verified by "
        "phone / from a professional email address",
        "References specifically asked about suitability to work with "
        "children / adults at risk and any disciplinary history",
        "Qualifications verified (original certificates seen)",
        "Professional registration checked where applicable (e.g. NMC, "
        "Social Work England)",
        "Overseas criminal record check / certificate of good conduct "
        "obtained where the candidate has lived or worked abroad",
        "Health declaration: physically and mentally fit for the role "
        "(with reasonable adjustments considered)",
        "Conditional offer letter issued (subject to checks)",
        "All checks completed BEFORE unsupervised work commenced",
        "Single central record / staff register updated",
        "Probation period, supervision schedule and induction (Care "
        "Certificate where new to care) arranged",
    ], first_col="Stage", extra=("Date completed", "Initials"), extra_w=2.8)
    h1(doc, "Sign-off")
    field_table(doc, [
        ("Any information of concern and how it was risk assessed",
         "e.g. disclosed convictions, reference concerns, gaps"),
    ], row_h=1.2)
    sig_table(doc, ["Recruiting manager", "Registered manager"])
    return doc


def t_audit_schedule_annual():
    doc = new_doc("Annual Audit Schedule", landscape=True)
    guidance(doc, "Plan your internal audit cycle for the year: mark the "
                  "planned month(s) for each audit area, then record "
                  "completion dates. Feed findings into your action plan "
                  "and governance reviews (Reg 45 children's homes / Reg 17 "
                  "good governance).")
    field_table(doc, ["Service", "Year", "Schedule owner"],
                header="Schedule details", label_w=5.0, total_w=LANDSCAPE_W)
    h1(doc, "12-month planner")
    guidance(doc, "Suggested frequencies are examples — set your own and "
                  "mark planned months with 'P', completed with the date.")
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug",
              "Sep", "Oct", "Nov", "Dec"]
    areas = [
        "Medication / MAR charts (monthly)",
        "Care / placement plans and reviews (monthly sample)",
        "Daily records quality (monthly sample)",
        "Incidents, accidents and near misses (monthly trend review)",
        "Safeguarding log and referrals (quarterly)",
        "Complaints and compliments (quarterly)",
        "Health and safety / premises walkaround (monthly)",
        "Fire safety: alarms, drills, PEEPs, equipment (monthly / "
        "quarterly)",
        "Infection prevention and control (quarterly)",
        "Kitchen / food hygiene and fridge temps (monthly)",
        "Finance: personal monies / petty cash (monthly)",
        "Supervision and appraisal compliance (quarterly)",
        "Training matrix compliance (quarterly)",
        "Recruitment files / single central record (quarterly sample)",
        "Rotas and dependency / staffing levels (quarterly)",
        "Maintenance and vehicle checks (monthly)",
        "Night-time provision spot check (quarterly, unannounced)",
        "Data protection / records security (annual)",
        "Business continuity plan test (annual)",
    ]
    widths = [7.5] + [1.3] * 12 + [3.0]
    grid_table(doc, ["Audit area"] + months + ["Lead"],
               widths, blank_rows=3, row_h=0.7, font_size=8,
               prefill=[[a] + [""] * 13 for a in areas])
    for r in doc.tables[-1].rows[1:]:
        row_height(r, 0.65)
    h1(doc, "Key")
    para(doc, "P = planned   •   Enter completion date when done   •   "
              "R = rescheduled (note new month)   •   Findings and actions "
              "are logged on the audit action plan, reviewed monthly by "
              "the manager.", size=9)
    return doc


def t_risk_assessment_5x5():
    doc = new_doc("Risk Assessment (5x5 Matrix)")
    guidance(doc, "General risk assessment template using a 5x5 likelihood "
                  "x severity matrix. Suitable for activities, environments "
                  "and individual risks. Review at the stated date, after "
                  "any incident, and when anything significant changes.")
    field_table(doc, [
        ("Assessment of", "Activity / task / environment / person"),
        "Reference number", "Assessor (name, role)", "Date of assessment",
        "Review date", "People who could be affected",
    ], header="Assessment details", label_w=6.0)

    h1(doc, "Scoring key")
    grid_table(doc, ["Score", "Likelihood", "Severity (consequence)"],
               [1.6, 7.6, 8.2], blank_rows=0, prefill=[
        ["1", "Rare — very unlikely to happen",
         "Negligible — no injury or very minor, no treatment"],
        ["2", "Unlikely — could happen but not expected",
         "Minor — first aid treatment only"],
        ["3", "Possible — might happen occasionally",
         "Moderate — medical treatment, up to 7 days off"],
        ["4", "Likely — will probably happen at some point",
         "Major — serious injury, RIDDOR reportable, hospitalisation"],
        ["5", "Almost certain — expected to happen, possibly often",
         "Catastrophic — death or permanent disability"],
    ])

    h2(doc, "Risk matrix (likelihood x severity = risk score)")
    matrix = doc.add_table(rows=6, cols=6)
    matrix.style = "Table Grid"
    fixed_layout(matrix)
    band_fill = {"low": "C6E0B4", "med": "FFE699", "high": "F4B183",
                 "vhigh": "FF7C80"}

    def band(score):
        if score <= 4:
            return "low"
        if score <= 9:
            return "med"
        if score <= 15:
            return "high"
        return "vhigh"

    cell_text(matrix.rows[0].cells[0], "Likelihood ↓  Severity →",
              bold=True, size=8, fill=DARK_FILL, color=WHITE)
    for s in range(1, 6):
        cell_text(matrix.rows[0].cells[s], str(s), bold=True, color=WHITE,
                  fill=DARK_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    for l in range(1, 6):
        row = matrix.rows[l]
        cell_text(row.cells[0], str(l), bold=True, color=WHITE,
                  fill=DARK_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
        for s in range(1, 6):
            score = l * s
            cell_text(row.cells[s], str(score), bold=True,
                      fill=band_fill[band(score)],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in matrix.rows:
        row_height(row, 0.65)
        row.cells[0].width = Cm(4.4)
        for c in row.cells[1:]:
            c.width = Cm(2.6)
    doc.add_paragraph()
    grid_table(doc, ["Score band", "Rating", "Required action"],
               [3.0, 3.0, 11.4], blank_rows=0, prefill=[
        ["1–4", "Low", "Acceptable — manage by routine procedures; monitor."],
        ["5–9", "Medium", "Action required to reduce risk; implement "
                          "within a defined timescale."],
        ["10–15", "High", "Urgent action required; senior management "
                          "attention; restrict activity until controls in "
                          "place."],
        ["16–25", "Very high", "Stop the activity immediately; do not "
                               "proceed until the risk is reduced."],
    ])

    h1(doc, "Hazard assessment")
    grid_table(doc, ["Hazard", "Who might be harmed and how",
                     "Existing controls", "L", "S", "Score",
                     "Further action required (owner, by when)",
                     "Residual L x S"],
               [2.6, 3.0, 3.6, 0.9, 0.9, 1.3, 3.6, 1.5],
               blank_rows=6, row_h=1.5, font_size=8.5)

    h1(doc, "Sign-off and review")
    sig_table(doc, ["Assessor", "Manager (approval)"])
    grid_table(doc, ["Review date", "Reviewed by",
                     "Changes made / still valid?", "Signature"],
               [2.6, 3.6, 7.6, 3.6], blank_rows=3, row_h=0.7)
    return doc


# ------------------------------------------------------------------- XLSX ---

def build_dols_tracker(path):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = Workbook()
    ws = wb.active
    ws.title = "DoLS Tracker"

    headers = [
        ("Resident name", 22), ("Room", 8), ("Date of birth", 12),
        ("Date applied", 12), ("Urgent authorisation? (Y/N)", 14),
        ("Urgent expiry (7 days, extendable to 14)", 16),
        ("Standard application sent", 14),
        ("Supervisory body (local authority)", 24), ("Status", 24),
        ("Date granted", 12), ("Expiry date", 12),
        ("Conditions attached", 30), ("Actions on conditions", 30),
        ("Next internal review date", 14), ("RPR name and contact", 24),
        ("CQC notified (application & outcome)", 16), ("Notes", 34),
    ]
    header_fill = PatternFill("solid", fgColor="1F4E5F")
    header_font = Font(color="FFFFFF", bold=True, size=10)
    wrap = Alignment(wrap_text=True, vertical="top")
    for col, (title, width) in enumerate(headers, start=1):
        c = ws.cell(row=1, column=col, value=title)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(wrap_text=True, vertical="center")
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.row_dimensions[1].height = 42
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = "A1:Q1"
    for row in range(2, 42):
        for col in range(1, len(headers) + 1):
            ws.cell(row=row, column=col).alignment = wrap

    statuses = [
        ("Not yet applied", "Deprivation identified — application being "
                            "prepared. Do not delay."),
        ("Urgent in place", "Provider has granted itself an urgent "
                            "authorisation (max 7 days, one extension to 14) "
                            "and MUST have sent a standard application at "
                            "the same time."),
        ("Standard application pending", "Standard application with the "
                                         "supervisory body awaiting "
                                         "assessment / decision."),
        ("Granted", "Standard authorisation in place — diarise expiry and "
                    "apply for renewal in good time."),
        ("Granted with conditions", "Authorisation in place with conditions "
                                    "— log each condition and evidence "
                                    "compliance in the Actions column."),
        ("Refused", "Assessment concluded criteria not met — review the "
                    "care plan: is the person still deprived of their "
                    "liberty? Seek advice."),
        ("Expired — renewal applied", "Authorisation expired or expiring; "
                                      "renewal submitted."),
        ("No longer required", "Person moved, deceased, or regained "
                               "capacity / arrangements changed — notify "
                               "the supervisory body and CQC as required."),
    ]
    dv_status = DataValidation(
        type="list",
        formula1='"' + ",".join(s for s, _ in statuses) + '"',
        allow_blank=True, showDropDown=False)
    dv_status.error = "Choose a status from the list (see Status Key sheet)."
    dv_status.errorTitle = "Invalid status"
    ws.add_data_validation(dv_status)
    dv_status.add("I2:I200")
    dv_yn = DataValidation(type="list", formula1='"Y,N"', allow_blank=True,
                           showDropDown=False)
    ws.add_data_validation(dv_yn)
    dv_yn.add("E2:E200")

    key = wb.create_sheet("Status Key")
    key.cell(row=1, column=1, value="DoLS Application Tracker — Status Key")
    key.cell(row=1, column=1).font = Font(bold=True, size=13,
                                          color="1F4E5F")
    key.cell(row=2, column=1, value=INTRO).font = Font(italic=True, size=9,
                                                       color="595959")
    kh1 = key.cell(row=4, column=1, value="Status")
    kh2 = key.cell(row=4, column=2, value="Meaning / action required")
    for c in (kh1, kh2):
        c.fill = header_fill
        c.font = header_font
    key.column_dimensions["A"].width = 30
    key.column_dimensions["B"].width = 90
    r = 5
    for status, meaning in statuses:
        key.cell(row=r, column=1, value=status).alignment = wrap
        key.cell(row=r, column=2, value=meaning).alignment = wrap
        r += 1
    r += 1
    notes = [
        "Reminders:",
        "• An urgent authorisation lasts a maximum of 7 days and can be "
        "extended once by the supervisory body to 14 days.",
        "• A standard application must accompany every urgent "
        "authorisation.",
        "• Notify CQC of the application and of its outcome (statutory "
        "notification).",
        "• Review this tracker at least monthly and 8 weeks before any "
        "expiry date.",
        "• DoLS applies to care homes and hospitals; for supported living, "
        "deprivations must be authorised by the Court of Protection.",
    ]
    for n in notes:
        key.cell(row=r, column=1, value=n).alignment = wrap
        key.merge_cells(start_row=r, start_column=1, end_row=r,
                        end_column=2)
        r += 1
    r += 1
    key.cell(row=r, column=1, value=FOOTER).font = Font(italic=True, size=9,
                                                        color="595959")
    wb.save(path)


# ---------------------------------------------------------------- registry ---

TEMPLATES = [
    # OFSTED
    dict(slug="placement-plan", title="Placement Plan", regulator="ofsted",
         category="care-planning", builder=t_placement_plan,
         description="A complete placement plan covering the child's "
                     "details, legal status, delegated authority and "
                     "consents, health, education, identity, family time, "
                     "risks and daily living. Ready to complete within the "
                     "first days of a placement and review alongside the "
                     "statutory care plan.",
         supports="Reg 5, Care Planning, Placement and Case Review "
                  "(England) Regulations 2010; Reg 14, Children's Homes "
                  "(England) Regulations 2015"),
    dict(slug="behaviour-support-plan", title="Behaviour Support Plan (PBS)",
         regulator="ofsted", category="care-planning",
         builder=t_behaviour_support_plan,
         description="A positive behaviour support plan structured around "
                     "understanding the behaviour and primary, secondary "
                     "and reactive strategies, with post-incident repair. "
                     "Includes a ready-to-print ABC (Antecedent–Behaviour–"
                     "Consequence) recording sheet.",
         supports="Regs 11, 20 & 35, Children's Homes (England) "
                  "Regulations 2015"),
    dict(slug="physical-intervention-record",
         title="Physical Intervention Record", regulator="ofsted",
         category="safeguarding", builder=t_physical_intervention_record,
         description="A restraint record capturing everyone involved, "
                     "de-escalation attempted, the technique and duration, "
                     "injuries, effectiveness, the child's views and "
                     "management sign-off. Built around the 48-hour and "
                     "5-day requirements of Regulation 35.",
         supports="Reg 35, Children's Homes (England) Regulations 2015"),
    dict(slug="missing-from-care-record",
         title="Missing from Care Record & Return Home Interview",
         regulator="ofsted", category="safeguarding",
         builder=t_missing_from_care_record,
         description="A two-part record for missing episodes: risk "
                     "information, police reference, a timeline of actions "
                     "and return details, followed by an independent return "
                     "home interview form with suggested questions and an "
                     "action plan.",
         supports="Statutory guidance on children who run away or go "
                  "missing from home or care (DfE, 2014); Reg 40, "
                  "Children's Homes (England) Regulations 2015"),
    dict(slug="key-work-session-log", title="Key Work Session Log",
         regulator="ofsted", category="recording",
         builder=t_key_work_session_log,
         description="A structured log for one-to-one key work sessions "
                     "linking each session to care plan objectives, "
                     "recording the child's views in their own words and "
                     "tracking agreed actions to the next session.",
         supports="Quality Standards (Regs 6–14), Children's Homes "
                  "(England) Regulations 2015"),
    dict(slug="reg-44-visit-pack",
         title="Reg 44 Independent Visitor Preparation Pack",
         regulator="ofsted", category="quality-audit",
         builder=t_reg_44_visit_pack,
         description="Everything the independent person needs for a "
                     "monthly Regulation 44 visit: a document sampling "
                     "checklist, suggested questions for children, staff "
                     "and the manager, observation prompts and a report "
                     "structure focused on the two required opinions.",
         supports="Reg 44, Children's Homes (England) Regulations 2015"),
    dict(slug="reg-45-quality-review", title="Reg 45 Quality of Care Review",
         regulator="ofsted", category="quality-audit",
         builder=t_reg_45_quality_review,
         description="A six-monthly quality of care review structured "
                     "around the Reg 45(2) matters: quality of care data "
                     "and Quality Standards self-assessment, children's "
                     "feedback and opinions, complaint actions and an "
                     "improvement plan.",
         supports="Reg 45, Children's Homes (England) Regulations 2015"),
    dict(slug="daily-log-template", title="Daily Log", regulator="ofsted",
         category="recording", builder=t_daily_log_template,
         description="A daily recording template with built-in prompts for "
                     "factual, non-judgemental, child-focused recording "
                     "across the whole day, including positives as well as "
                     "concerns. Written on the basis that the child may "
                     "read it.",
         supports="Reg 36 & Schedule 3, Children's Homes (England) "
                  "Regulations 2015"),
    dict(slug="handover-sheet", title="Shift Handover Sheet",
         regulator="ofsted", category="recording", builder=t_handover_sheet,
         description="A shift-change handover covering each child's day, "
                     "medication, appointments, incidents, safeguarding "
                     "updates, money, vehicles and outstanding tasks, "
                     "signed by both the outgoing and incoming senior.",
         supports="Reg 13 (leadership and management), Children's Homes "
                  "(England) Regulations 2015"),
    dict(slug="friendship-associates-log", title="Friendship & Associates Log",
         regulator="ofsted", category="safeguarding",
         builder=t_friendship_associates_log,
         description="A running log of each child's friends and associates "
                     "on and offline, with prompts for exploitation "
                     "warning signs and a clear route into safeguarding "
                     "action. Supports missing episode responses and "
                     "contextual safeguarding.",
         supports="Reg 12 (protection of children), Children's Homes "
                  "(England) Regulations 2015"),
    dict(slug="young-persons-guide-skeleton",
         title="Young Person's Guide (Skeleton)", regulator="ofsted",
         category="care-planning", builder=t_young_persons_guide_skeleton,
         description="A child-friendly skeleton for your children's guide, "
                     "with thirteen ready-made headings from 'Welcome' to "
                     "'Important numbers' and guidance on what to write in "
                     "each. Adapt the language to the ages and needs of "
                     "the children you care for.",
         supports="Children's guide requirements, Children's Homes "
                  "(England) Regulations 2015 and the Guide to the "
                  "Regulations"),
    # CQC
    dict(slug="person-centred-care-plan", title="Person-Centred Care Plan",
         regulator="cqc", category="care-planning",
         builder=t_person_centred_care_plan,
         description="A full person-centred care plan: about me, "
                     "communication, consent and capacity, and twelve care "
                     "domains each recording what the person can do, the "
                     "support they want and the outcome being worked "
                     "towards, with a review record.",
         supports="Regs 9 & 11, Health and Social Care Act 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="one-page-profile", title="One-Page Profile", regulator="cqc",
         category="care-planning", builder=t_one_page_profile,
         description="The classic one-page profile: what people appreciate "
                     "about me, what is important to me, and how best to "
                     "support me — with hints that push for specific, "
                     "personal detail a stranger could act on.",
         supports="Regs 9 & 10, Health and Social Care Act 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="mca-capacity-assessment", title="MCA Capacity Assessment",
         regulator="cqc", category="care-planning",
         builder=t_mca_capacity_assessment,
         description="A decision-specific mental capacity assessment "
                     "working through the two-stage test: the diagnostic "
                     "test and the four functional elements (understand, "
                     "retain, use or weigh, communicate), with the "
                     "supported decision-making steps evidenced first.",
         supports="Mental Capacity Act 2005 ss.1–3; Reg 11, HSCA 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="best-interests-decision-record",
         title="Best Interests Decision Record", regulator="cqc",
         category="care-planning", builder=t_best_interests_decision_record,
         description="A section 4 best interests record with the full "
                     "statutory checklist, a consultation table, an "
                     "options comparison including less restrictive "
                     "alternatives, and the decision with rationale and "
                     "review arrangements.",
         supports="Mental Capacity Act 2005 s.4"),
    dict(slug="dols-application-tracker", title="DoLS Application Tracker",
         regulator="cqc", category="safeguarding", fmt="xlsx",
         description="A spreadsheet tracker for DoLS applications: urgent "
                     "and standard authorisations, supervisory body, "
                     "status, expiry and review dates, conditions and RPR "
                     "details, with a status key sheet and dropdown "
                     "validation built in.",
         supports="Mental Capacity Act 2005 Schedule A1 (DoLS); Reg 13, "
                  "HSCA 2008 (Regulated Activities) Regulations 2014"),
    dict(slug="mar-audit-tool", title="MAR Audit Tool", regulator="cqc",
         category="medication", builder=t_mar_audit_tool,
         description="A twenty-question MAR chart audit with Yes / No / "
                     "N/A columns covering signatures, gaps, PRN "
                     "protocols, controlled drugs, storage temperatures "
                     "and staff competency, feeding into an action plan.",
         supports="Reg 12(2)(g) (safe management of medicines), HSCA 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="medication-error-record", title="Medication Error Record",
         regulator="cqc", category="medication",
         builder=t_medication_error_record,
         description="A no-blame medication error record: what happened, "
                     "immediate safety checks, GP / NHS 111 advice, "
                     "notifications including CQC and safeguarding, and a "
                     "root-cause reflection with actions to prevent "
                     "recurrence.",
         supports="Reg 12, HSCA 2008 (Regulated Activities) Regulations "
                  "2014; Reg 18, CQC (Registration) Regulations 2009"),
    dict(slug="incident-accident-form", title="Incident / Accident Form",
         regulator="cqc", category="recording",
         builder=t_incident_accident_form,
         description="An incident and accident form for residents, staff "
                     "and visitors with injuries and treatment, a "
                     "structured RIDDOR reportability check, notification "
                     "checklist and manager review for learning.",
         supports="Regs 12 & 17, HSCA 2008 (Regulated Activities) "
                  "Regulations 2014; RIDDOR 2013"),
    dict(slug="complaints-log", title="Complaints Log", regulator="cqc",
         category="quality-audit", builder=t_complaints_log,
         description="A single oversight log for all complaints — "
                     "including verbal ones — tracking acknowledgement, "
                     "investigation, outcome, response and learning, with "
                     "recommended timescales and escalation routes "
                     "included.",
         supports="Reg 16 (receiving and acting on complaints), HSCA 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="duty-of-candour-record", title="Duty of Candour Record",
         regulator="cqc", category="quality-audit",
         builder=t_duty_of_candour_record,
         description="A step-by-step Regulation 20 record: the notifiable "
                     "safety incident, the verbal notification and "
                     "apology, the written follow-up, sharing of "
                     "investigation results and the audit trail CQC "
                     "expects to see.",
         supports="Reg 20 (duty of candour), HSCA 2008 (Regulated "
                  "Activities) Regulations 2014"),
    dict(slug="statement-of-purpose-skeleton",
         title="Statement of Purpose (Skeleton)", regulator="cqc",
         category="quality-audit", builder=t_statement_of_purpose_skeleton,
         description="A ten-section skeleton covering the Schedule 3 "
                     "matters: provider details, aims and objectives, "
                     "regulated activities, needs met, locations, "
                     "staffing, the registered manager and contact "
                     "details, with guidance under each heading.",
         supports="Reg 12 & Schedule 3, CQC (Registration) Regulations "
                  "2009"),
    dict(slug="service-user-guide-skeleton",
         title="Service User Guide (Skeleton)", regulator="cqc",
         category="care-planning", builder=t_service_user_guide_skeleton,
         description="A plain-English service user guide skeleton with "
                     "twelve headings from 'Welcome' to 'Useful contacts', "
                     "each with guidance on content, adaptable to "
                     "residential or home care settings and accessible "
                     "formats.",
         supports="Regs 9 & 17, HSCA 2008 (Regulated Activities) "
                  "Regulations 2014; Accessible Information Standard"),
    # SHARED
    dict(slug="supervision-record-staff", title="Supervision Record — Staff",
         regulator="both", category="hr-staffing",
         builder=t_supervision_record_staff,
         description="A one-to-one supervision record for care staff "
                     "covering wellbeing, caseload and key work, "
                     "safeguarding, performance, development and agreed "
                     "actions, with review of the previous session's "
                     "actions built in.",
         supports="Reg 33, Children's Homes (England) Regulations 2015; "
                  "Reg 18(2)(a), HSCA 2008 (Regulated Activities) "
                  "Regulations 2014"),
    dict(slug="supervision-record-manager",
         title="Supervision Record — Manager", regulator="both",
         category="hr-staffing", builder=t_supervision_record_manager,
         description="A supervision record for the registered manager, "
                     "led by the responsible or nominated individual: "
                     "safeguarding and incident oversight, quality and "
                     "compliance, staffing, resources and the manager's "
                     "own wellbeing and development.",
         supports="Reg 33, Children's Homes (England) Regulations 2015; "
                  "Reg 18(2)(a), HSCA 2008 (Regulated Activities) "
                  "Regulations 2014"),
    dict(slug="team-meeting-agenda-minutes",
         title="Team Meeting Agenda & Minutes", regulator="both",
         category="hr-staffing", builder=t_team_meeting_agenda_minutes,
         description="A combined agenda and minutes template with a "
                     "twelve-item standing agenda, an action log, and a "
                     "read-and-understood sign sheet for staff who missed "
                     "the meeting.",
         supports="Reg 13, Children's Homes (England) Regulations 2015; "
                  "Reg 17, HSCA 2008 (Regulated Activities) Regulations "
                  "2014"),
    dict(slug="reflective-practice-log", title="Reflective Practice Log",
         regulator="both", category="hr-staffing",
         builder=t_reflective_practice_log,
         description="A reflection template working through all six "
                     "stages of Gibbs' reflective cycle with prompts at "
                     "each stage, linking the reflection into supervision "
                     "and the staff member's development plan.",
         supports="Care Certificate Standard 2; Reg 18, HSCA 2008 "
                  "(Regulated Activities) Regulations 2014"),
    dict(slug="induction-checklist-care-certificate",
         title="Induction Checklist (Care Certificate)", regulator="both",
         category="hr-staffing",
         builder=t_induction_checklist_care_certificate,
         description="A two-part induction checklist: first-week "
                     "essentials from fire safety to key policies, then "
                     "sign-off against each of the 15 Care Certificate "
                     "standards with assessor initials and a 12-week "
                     "target.",
         supports="The Care Certificate (Skills for Care / Health "
                  "Education England); Reg 18 & 19, HSCA 2008 (Regulated "
                  "Activities) Regulations 2014; Reg 33, CHR 2015"),
    dict(slug="safer-recruitment-checklist",
         title="Safer Recruitment Checklist", regulator="both",
         category="hr-staffing", builder=t_safer_recruitment_checklist,
         description="A start-to-finish safer recruitment checklist for "
                     "each appointment — advert wording, application and "
                     "gap checking, enhanced DBS with barred list, "
                     "verified references, overseas checks and induction — "
                     "filed as evidence in the staff record.",
         supports="Schedule 2, Children's Homes (England) Regulations "
                  "2015; Reg 19 & Schedule 3, HSCA 2008 (Regulated "
                  "Activities) Regulations 2014"),
    dict(slug="audit-schedule-annual", title="Annual Audit Schedule",
         regulator="both", category="quality-audit",
         builder=t_audit_schedule_annual,
         description="A landscape 12-month audit planner pre-loaded with "
                     "nineteen audit areas and suggested frequencies, from "
                     "MAR charts to unannounced night visits, with a "
                     "simple planned / completed marking key.",
         supports="Reg 45, Children's Homes (England) Regulations 2015; "
                  "Reg 17 (good governance), HSCA 2008 (Regulated "
                  "Activities) Regulations 2014"),
    dict(slug="risk-assessment-5x5", title="Risk Assessment (5x5 Matrix)",
         regulator="both", category="safeguarding",
         builder=t_risk_assessment_5x5,
         description="A general risk assessment template with a colour "
                     "5x5 likelihood-by-severity matrix, full scoring key "
                     "and rating bands, hazard table with residual "
                     "scoring, and a review record.",
         supports="Management of Health and Safety at Work Regulations "
                  "1999; Reg 12, CHR 2015; Reg 12, HSCA 2008 (Regulated "
                  "Activities) Regulations 2014"),
]


# ----------------------------------------------------------------- main ---

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    os.makedirs(SEED_DIR, exist_ok=True)

    generated = []
    for spec in TEMPLATES:
        slug = spec["slug"]
        fmt = spec.get("fmt", "docx")
        if fmt == "xlsx":
            path = os.path.join(OUT_DIR, slug + ".xlsx")
            build_dols_tracker(path)
        else:
            doc = spec["builder"]()
            path = finish(doc, slug)
        generated.append((slug, fmt, path))
        print("generated  {}".format(os.path.basename(path)))

    # Verify every file parses back
    from openpyxl import load_workbook
    for slug, fmt, path in generated:
        assert os.path.exists(path), "missing: " + path
        size = os.path.getsize(path)
        assert size > 5000, "suspiciously small ({} bytes): {}".format(
            size, path)
        if fmt == "xlsx":
            wb = load_workbook(path)
            assert "DoLS Tracker" in wb.sheetnames
            assert "Status Key" in wb.sheetnames
        else:
            d = Document(path)
            assert len(d.paragraphs) > 3, "too few paragraphs: " + path
            assert d.paragraphs[0].text, "no title: " + path
    print("verified   {} files parse back correctly".format(len(generated)))

    # Seed JSON
    seed = []
    for spec in TEMPLATES:
        fmt = spec.get("fmt", "docx")
        seed.append({
            "slug": spec["slug"],
            "title": spec["title"],
            "regulator": spec["regulator"],
            "category": spec["category"],
            "format": fmt,
            "filename": spec["slug"] + "." + fmt,
            "description": spec["description"],
            "supports": spec["supports"],
        })
    seed_path = os.path.join(SEED_DIR, "templates.json")
    with open(seed_path, "w", encoding="utf-8") as f:
        json.dump(seed, f, indent=2, ensure_ascii=False)
        f.write("\n")
    with open(seed_path, encoding="utf-8") as f:
        loaded = json.load(f)
    assert len(loaded) == len(TEMPLATES)
    for entry in loaded:
        assert os.path.exists(os.path.join(OUT_DIR, entry["filename"])), (
            "seed filename has no generated file: " + entry["filename"])
    print("seed file  {} ({} entries)".format(seed_path, len(loaded)))
    print("DONE: {} template files generated in {}".format(
        len(generated), OUT_DIR))


if __name__ == "__main__":
    sys.exit(main())
